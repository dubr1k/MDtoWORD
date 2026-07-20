from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Any, cast
from urllib.request import urlopen

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.styles.style import ParagraphStyle
from markdown_it import MarkdownIt
from mdit_py_plugins.amsmath import amsmath_plugin
from mdit_py_plugins.dollarmath import dollarmath_plugin
from mdit_py_plugins.footnote import footnote_plugin

from latex_omml import UnsupportedLatexError, latex_to_omml


_TASK_PREFIX = re.compile(r"^\[([ xX])\]\s+")
_CYRILLIC = re.compile(r"[Ѐ-ӿ]")
_BLACK = RGBColor(0, 0, 0)

# ``latex_omml`` parses these environments itself, so they are passed through
# with their ``\begin``/``\end`` wrapper intact.
_MATRIX_ENVIRONMENTS = frozenset(
    {"matrix", "pmatrix", "bmatrix", "Bmatrix", "vmatrix", "Vmatrix"}
)
# The amsmath plugin hands us the environment complete with its wrapper, e.g.
# "\begin{align}\na &= b \\\\\nc &= d\n\end{align}". ``alignat`` adds a column
# count right after the opening tag, hence the optional brace group.
_AMSMATH_WRAPPER = re.compile(
    r"^\\begin\{(?P<environment>[A-Za-z]+)\*?\}(?:\{[^{}]*\})?"
    r"(?P<body>.*)"
    r"\\end\{(?P=environment)\*?\}$",
    re.DOTALL,
)
_LINE_BREAK = re.compile(r"\\\\")
# An unescaped "&" is amsmath column alignment; "\&" is a literal ampersand.
_ALIGNMENT_MARKER = re.compile(r"(?<!\\)&")


class GfmDocxRenderer:
    """Render a GFM token stream into a Word document."""

    def __init__(self, font_name: str, font_size: Pt):
        self.font_name = font_name
        self.font_size = font_size
        self.document: DocumentType
        self.warnings: list[str]
        self._paragraph: Any
        self._list_stack: list[str]
        self._quote_depth: int
        self._table_rows: list[list[str]] | None
        self._table_row: list[str] | None
        self._table_cell: list[str] | None
        self._table_alignments: list[str | None]
        self._table_header: bool
        self._footnote_depth: int

    def render(
        self, markdown: str, source_path: Path | None = None
    ) -> tuple[DocumentType, list[str]]:
        self.document = Document()
        self.warnings = []
        self._paragraph = None
        self._list_stack = []
        self._quote_depth = 0
        self._table_rows = None
        self._table_row = None
        self._table_cell = None
        self._table_header = False
        self._table_alignments = []
        self._footnote_depth = 0
        self._configure_document()

        parser = (
            MarkdownIt("js-default", {"breaks": True, "html": False, "linkify": True})
            .enable("linkify")
            .use(footnote_plugin)
            .use(dollarmath_plugin, allow_digits=False, allow_blank_lines=False)
            .use(amsmath_plugin)
        )
        for token in parser.parse(markdown):
            self._render_block(token, source_path)

        return self.document, self.warnings

    def _configure_document(self) -> None:
        style = cast(ParagraphStyle, self.document.styles["Normal"])
        style.font.name = self.font_name
        style.font.size = self.font_size
        style.font.color.rgb = _BLACK
        for level in range(1, 10):
            try:
                heading = cast(ParagraphStyle, self.document.styles[f"Heading {level}"])
            except KeyError:
                continue
            heading.font.color.rgb = _BLACK
            heading.font.name = self.font_name
        try:
            quote = cast(ParagraphStyle, self.document.styles["Quote"])
        except KeyError:
            pass
        else:
            quote.font.color.rgb = _BLACK

    def _render_block(self, token: Any, source_path: Path | None) -> None:
        token_type = token.type

        if token_type == "heading_open":
            level = min(int(token.tag[1:]), 9)
            self._paragraph = self.document.add_paragraph(style=f"Heading {level}")
            return
        if token_type == "paragraph_open":
            self._paragraph = self._new_paragraph()
            return
        if token_type == "inline":
            self._render_inline(token.children or [], source_path, token.content)
            return
        if token_type in {"heading_close", "paragraph_close"}:
            self._paragraph = None
            return
        if token_type == "blockquote_open":
            self._quote_depth += 1
            return
        if token_type == "blockquote_close":
            self._quote_depth -= 1
            return
        if token_type in {"bullet_list_open", "ordered_list_open"}:
            self._list_stack.append(token_type)
            return
        if token_type in {"bullet_list_close", "ordered_list_close"}:
            self._list_stack.pop()
            return
        if token_type in {"list_item_open", "list_item_close"}:
            return
        if token_type == "amsmath":
            self._render_amsmath(token)
            return
        if token_type == "math_block":
            self._render_math(token.content, display=True)
            return
        if token_type == "math_block_label":
            self._render_math(token.content, display=True)
            self._append_equation_label(token.info)
            return
        if token_type in {"fence", "code_block"}:
            self._render_code_block(token)
            return
        if token_type == "hr":
            self._add_thematic_break()
            return
        if token_type == "table_open":
            self._table_rows = []
            self._table_alignments = []
            return
        if token_type == "thead_open":
            self._table_header = True
            return
        if token_type == "thead_close":
            self._table_header = False
            return
        if token_type == "tr_open":
            self._table_row = []
            return
        if token_type in {"th_open", "td_open"}:
            self._table_cell = []
            if self._table_header:
                style_attr = token.attrGet("style") or ""
                if "right" in style_attr:
                    self._table_alignments.append("right")
                elif "center" in style_attr:
                    self._table_alignments.append("center")
                else:
                    self._table_alignments.append(None)
            return
        if token_type in {"th_close", "td_close"}:
            if self._table_row is not None and self._table_cell is not None:
                self._table_row.append("".join(self._table_cell))
            self._table_cell = None
            return
        if token_type == "tr_close":
            if self._table_rows is not None and self._table_row is not None:
                self._table_rows.append(self._table_row)
            self._table_row = None
            return
        if token_type == "table_close":
            self._finish_table()
            return
        if token_type == "footnote_block_open":
            self.document.add_heading("Footnotes", level=2)
            self._footnote_depth += 1
            return
        if token_type == "footnote_block_close":
            self._footnote_depth -= 1
            return
        if token_type == "footnote_open":
            label = token.meta["label"]
            self._paragraph = self.document.add_paragraph(style="List Number")
            self._paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._paragraph.add_run(f"[{label}] ")
            return
        if token_type == "footnote_close":
            self._paragraph = None

    def _new_paragraph(self):
        if self._list_stack:
            style = "List Number" if self._list_stack[-1] == "ordered_list_open" else "List Bullet"
            paragraph = self.document.add_paragraph(style=style)
            paragraph.paragraph_format.left_indent = Pt(18 * (len(self._list_stack) - 1))
        elif self._quote_depth:
            paragraph = self.document.add_paragraph(style="Quote")
        else:
            paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return paragraph

    def _render_inline(
        self, children: list[Any], source_path: Path | None, source_content: str
    ) -> None:
        if self._table_cell is not None:
            self._table_cell.append(self._inline_text(children))
            return
        if self._paragraph is None:
            self._paragraph = self._new_paragraph()

        task_match = _TASK_PREFIX.match(source_content) if self._list_stack else None
        if task_match:
            self._paragraph.add_run("☒ " if task_match.group(1).lower() == "x" else "☐ ")
            self._skip_task_prefix(children, task_match.group(0))

        formatting = {"bold": False, "italic": False, "strike": False, "code": False}
        link_target: str | None = None
        for token in children:
            token_type = token.type
            if token_type == "text":
                self._append_text(token.content, formatting, link_target)
            elif token_type in {"softbreak", "hardbreak"}:
                self._paragraph.add_run().add_break()
            elif token_type == "code_inline":
                self._append_text(token.content, {**formatting, "code": True}, link_target)
            elif token_type == "em_open":
                formatting["italic"] = True
            elif token_type == "em_close":
                formatting["italic"] = False
            elif token_type == "strong_open":
                formatting["bold"] = True
            elif token_type == "strong_close":
                formatting["bold"] = False
            elif token_type == "s_open":
                formatting["strike"] = True
            elif token_type == "s_close":
                formatting["strike"] = False
            elif token_type == "link_open":
                link_target = token.attrGet("href")
            elif token_type == "link_close":
                link_target = None
            elif token_type == "image":
                self._append_image(token, source_path)
            elif token_type == "footnote_ref":
                self._append_text(f"[{token.meta['label']}]", formatting, None)
            elif token_type in {"math_inline", "math_inline_double"}:
                self._render_math(
                    token.content, display=False, markup=token.markup or "$"
                )
            else:
                self._append_text(token.content, formatting, link_target)

    @staticmethod
    def _skip_task_prefix(children: list[Any], prefix: str) -> None:
        for token in children:
            if token.type == "text":
                token.content = token.content.removeprefix(prefix)
                return

    @staticmethod
    def _inline_text(children: list[Any]) -> str:
        return "".join(
            "\n" if token.type in {"softbreak", "hardbreak"} else token.content
            for token in children
            if token.type not in {"link_open", "link_close", "em_open", "em_close", "strong_open", "strong_close", "s_open", "s_close"}
        )

    def _append_text(
        self, text: str, formatting: dict[str, bool], link_target: str | None
    ) -> None:
        if not text:
            return
        if link_target:
            self._append_hyperlink(text, link_target, formatting)
            return
        run = self._paragraph.add_run(text)
        run.bold = formatting["bold"]
        run.italic = formatting["italic"]
        run.font.strike = formatting["strike"]
        run.font.name = "Courier New" if formatting["code"] else self.font_name
        run.font.size = Pt(10) if formatting["code"] else self.font_size

    def _append_hyperlink(
        self, text: str, target: str, formatting: dict[str, bool]
    ) -> None:
        relationship_id = self._paragraph.part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        properties.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(underline)
        if formatting["bold"]:
            properties.append(OxmlElement("w:b"))
        if formatting["italic"]:
            properties.append(OxmlElement("w:i"))
        if formatting["strike"]:
            properties.append(OxmlElement("w:strike"))
        run.append(properties)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        self._paragraph._p.append(hyperlink)

    def _append_image(self, token: Any, source_path: Path | None) -> None:
        target = token.attrGet("src") or ""
        alt_text = token.content or "image"
        try:
            if target.startswith(("http://", "https://")):
                with urlopen(target, timeout=10) as response:
                    image_data = BytesIO(response.read())
                self._paragraph.add_run().add_picture(image_data)
                return

            image_path = Path(target)
            if not image_path.is_absolute() and source_path is not None:
                image_path = source_path.parent / image_path
            if not image_path.is_file():
                raise FileNotFoundError(target)
            self._paragraph.add_run().add_picture(str(image_path))
        except FileNotFoundError:
            self.warnings.append(f"Image not found: {target}")
            self._append_text(f"[{alt_text}]", {"bold": False, "italic": False, "strike": False, "code": False}, None)
        except Exception as error:
            self.warnings.append(f"Image could not be rendered: {target} ({error})")
            self._append_text(f"[{alt_text}]", {"bold": False, "italic": False, "strike": False, "code": False}, None)

    def _render_math(self, latex: str, display: bool, markup: str = "$") -> None:
        """Insert a real Word equation, falling back to verbatim text.

        Every path ends in either an equation or the untouched source plus a
        warning, so a formula can never silently vanish or be silently wrong.
        """
        formula = latex.strip("\n").strip()
        if not formula:
            if display:
                # Keep the empty block so an equation label still has a home.
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return
            # An empty inline span means the dollars were literal, as in
            # "x $ $ y"; put them back rather than swallowing them.
            if self._paragraph is None:
                self._paragraph = self._new_paragraph()
            self._append_text(
                f"{markup}{latex}{markup}",
                {"bold": False, "italic": False, "strike": False, "code": False},
                None,
            )
            return
        if not display and _CYRILLIC.search(formula):
            # Cyrillic between single dollars is far more often prose (shell
            # variables, price ranges) than a formula; keep it verbatim and
            # let _render_math_literal raise the "write \$" warning.
            self._render_math_literal(latex, display=False)
            return
        try:
            math_element = latex_to_omml(formula)
        except UnsupportedLatexError as error:
            self.warnings.append(f'Formula kept as text: "{formula}" ({error})')
            self._render_math_literal(latex, display)
            return
        self._place_math(math_element, display)

    def _place_math(self, math_element: Any, display: bool) -> None:
        if display:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph._p.append(math_element)
            return
        if self._paragraph is None:
            self._paragraph = self._new_paragraph()
        self._paragraph._p.append(math_element)

    def _render_amsmath(self, token: Any) -> None:
        r"""Render an ``amsmath`` environment as one or more Word equations.

        The plugin delivers the environment complete with its
        ``\begin{...}``/``\end{...}`` wrapper plus ``meta["environment"]``
        (star already stripped). Matrix families go to ``latex_omml``
        untouched because it parses them itself; everything else is unwrapped
        and split on ``\\`` into one centred equation per line, since OMML has
        no equivalent of amsmath's ``&`` column alignment. If any line fails
        to convert, the whole environment is kept verbatim so nothing of the
        source is lost to a partial rendering.
        """
        source = token.content
        environment = (token.meta or {}).get("environment", "")
        if environment in _MATRIX_ENVIRONMENTS:
            self._render_math(source, display=True)
            return

        match = _AMSMATH_WRAPPER.match(source.strip())
        body = match.group("body") if match else source
        lines = [
            _ALIGNMENT_MARKER.sub(" ", line).strip()
            for line in _LINE_BREAK.split(body)
        ]
        lines = [line for line in lines if line]
        if not lines:
            return

        elements = []
        for line in lines:
            try:
                elements.append(latex_to_omml(line))
            except UnsupportedLatexError as error:
                self.warnings.append(f'Formula kept as text: "{line}" ({error})')
                self._render_math_literal(source, display=True)
                return
        for element in elements:
            self._place_math(element, display=True)

    def _render_math_literal(self, latex: str, display: bool) -> None:
        """Write a formula as verbatim monospace text, preserving every character."""
        text = latex.strip("\n")
        if display:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
        else:
            if _CYRILLIC.search(latex):
                self.warnings.append(
                    f'Inline math "${latex}$" contains Cyrillic text and may be '
                    "ordinary prose rather than a formula; write a literal \"$\" "
                    'as "\\$".'
                )
            if self._paragraph is None:
                self._paragraph = self._new_paragraph()
            run = self._paragraph.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    def _append_equation_label(self, label: str) -> None:
        """Append an equation number after its display formula, tab-separated.

        Mirrors the LaTeX convention of a trailing equation number, e.g. the
        ``(1)`` in ``$$ x = 1 $$ (1)``. Rendered as ordinary body text (not
        monospace) in the same centred paragraph as the formula.
        """
        label = label.strip()
        if not label:
            return
        paragraph = self.document.paragraphs[-1]
        run = paragraph.add_run(f"\t({label})")
        run.font.name = self.font_name
        run.font.size = self.font_size

    def _render_code_block(self, token: Any) -> None:
        language = token.info.strip().split(maxsplit=1)[0] if token.info else ""
        if language:
            caption = self.document.add_paragraph()
            caption.add_run(language).italic = True
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(token.content.rstrip("\n"))
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    def _add_thematic_break(self) -> None:
        paragraph = self.document.add_paragraph()
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")
        borders.append(bottom)
        properties.append(borders)

    def _finish_table(self) -> None:
        if not self._table_rows:
            self._table_rows = None
            return
        columns = max(len(row) for row in self._table_rows)
        table = self.document.add_table(rows=len(self._table_rows), cols=columns)
        table.style = "Table Grid"
        self._apply_table_borders(table)
        alignments = {
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        for row_index, values in enumerate(self._table_rows):
            for column_index, value in enumerate(values):
                paragraph = table.cell(row_index, column_index).paragraphs[0]
                run = paragraph.add_run(value)
                if row_index == 0:
                    run.bold = True
                column_alignment = (
                    self._table_alignments[column_index]
                    if column_index < len(self._table_alignments)
                    else None
                )
                paragraph.alignment = alignments.get(
                    column_alignment or "", WD_ALIGN_PARAGRAPH.LEFT
                )
        self._table_rows = None
        self._table_alignments = []

    @staticmethod
    def _apply_table_borders(table: Any) -> None:
        """Write borders as direct formatting so every viewer renders them."""
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
            borders.append(element)
        table._tbl.tblPr.insert_element_before(
            borders,
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
            "w:tblCaption",
            "w:tblDescription",
            "w:tblPrChange",
        )
