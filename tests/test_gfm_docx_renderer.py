from pathlib import Path
import tempfile
import unittest
from unittest.mock import MagicMock, patch

from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor

from mdtoword.gfm_renderer import _MAX_REMOTE_IMAGE_BYTES, GfmDocxRenderer, _is_remote_target


_MATH_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

# The smallest well-formed PNG: an 8-byte signature, an IHDR chunk describing
# a single transparent pixel, an (empty-payload) IDAT chunk, and IEND. Real
# enough for python-docx's header parser to accept via ``add_picture``.
_MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\x05-\xb4\x00\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _urlopen_response(data: bytes) -> MagicMock:
    """Mock of ``urlopen(...)``'s return value, usable as a context manager."""
    response = MagicMock()
    response.read.return_value = data
    context = MagicMock()
    context.__enter__.return_value = response
    context.__exit__.return_value = False
    return context


class _FakeImageToken:
    """Minimal stand-in for markdown-it's image token.

    ``_append_image`` only ever reads ``.attrGet("src")`` and ``.content``,
    so a full parse round-trip is unnecessary -- and, for a UNC target,
    actively wrong: markdown-it percent-encodes any backslash in a link
    destination (``\\host`` becomes ``%5Chost``), which would silently
    launder away the exact string these tests need to exercise.
    """

    def __init__(self, src: str, alt: str = "diagram") -> None:
        self._src = src
        self.content = alt

    def attrGet(self, name: str) -> str | None:
        return self._src if name == "src" else None


def _renderer_ready_for_direct_image_append(renderer: GfmDocxRenderer) -> GfmDocxRenderer:
    """Put *renderer* in the state ``render()`` would, for calling
    ``_append_image`` directly with a hand-built token."""
    renderer.render("")
    renderer.warnings = []
    renderer._paragraph = renderer.document.add_paragraph()
    return renderer


# Every style GfmDocxRenderer applies to a paragraph somewhere in the
# renderer: Normal for plain body paragraphs, Heading 1-9 for headings,
# Quote for blockquotes, and List Bullet/List Number for list items.
_STYLES_THE_RENDERER_APPLIES = (
    ["Normal"] + [f"Heading {level}" for level in range(1, 10)]
    + ["Quote", "List Bullet", "List Number"]
)


def _equations(paragraph):
    """Every Word equation directly inside this paragraph."""
    return paragraph._p.findall(f"{_MATH_NS}oMath")


def _equation_text(equation):
    """All literal text inside one <m:oMath> element, concatenated."""
    return "".join(t.text or "" for t in equation.iter(f"{_MATH_NS}t"))


def _rfonts_attrs(style):
    """The raw ``w:rFonts`` attributes of ``style``, namespace-stripped.

    Reads the XML directly rather than going through python-docx's
    ``Font.name`` property, since that readback is exactly what hid the
    theme-font bug: python-docx reports the explicit ``ascii``/``hAnsi``
    value it just wrote even when a sibling ``*Theme`` attribute is what
    Word actually honours. Returns ``None`` if the style has no
    ``w:rFonts`` element at all.
    """
    style_element = style.element
    run_properties = style_element.find(qn("w:rPr"))
    if run_properties is None:
        return None
    rfonts = run_properties.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return {key.rsplit("}", 1)[-1]: value for key, value in rfonts.attrib.items()}


class GfmDocxRendererTests(unittest.TestCase):
    def setUp(self):
        self.renderer = GfmDocxRenderer("Arial", Pt(12))

    def test_renders_gfm_blocks_and_inline_formatting(self):
        document, warnings = self.renderer.render(
            "# Heading\n\n"
            "Paragraph with **bold**, *italic*, ~~struck~~, and `code`.\n\n"
            "> quoted text\n\n"
            "- [x] done\n"
            "  - nested\n\n"
            "| left | right |\n"
            "| :--- | ---: |\n"
            "| one | two |\n\n"
            "```python\n"
            "print('x')\n"
            "```"
        )

        self.assertEqual(document.paragraphs[0].style.name, "Heading 1")
        inline_paragraph = document.paragraphs[1]
        self.assertTrue(any(run.bold for run in inline_paragraph.runs))
        self.assertTrue(any(run.italic for run in inline_paragraph.runs))
        self.assertTrue(any(run.font.strike for run in inline_paragraph.runs))
        self.assertIn("quoted text", "\n".join(p.text for p in document.paragraphs))
        self.assertIn("☒ done", "\n".join(p.text for p in document.paragraphs))
        self.assertEqual(len(document.tables), 1)
        self.assertIn("print('x')", "\n".join(p.text for p in document.paragraphs))
        self.assertEqual(warnings, [])

    def test_renders_links_and_falls_back_when_local_image_is_missing(self):
        document, warnings = self.renderer.render(
            "[MDtoWord](https://example.com) and ![diagram](missing.png)",
            source_path=Path("/tmp/source.md"),
        )

        self.assertIn("MDtoWord", document.paragraphs[0].text)
        self.assertIn("[diagram]", document.paragraphs[0].text)
        self.assertEqual(warnings, ["Image not found: missing.png"])

    def test_remote_image_not_fetched_when_disallowed(self):
        renderer = GfmDocxRenderer("Arial", Pt(12), allow_remote_images=False)

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            document, warnings = renderer.render(
                "![diagram](https://example.invalid/x.png)"
            )

        mock_urlopen.assert_not_called()
        self.assertEqual(
            warnings,
            [
                "Remote image not fetched: https://example.invalid/x.png "
                "(network access is disabled; pass fetch_remote_images=true to allow it)"
            ],
        )
        self.assertIn("[diagram]", document.paragraphs[0].text)

    def test_remote_image_fetched_by_default(self):
        renderer = GfmDocxRenderer("Arial", Pt(12))

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            mock_urlopen.return_value = _urlopen_response(_MINIMAL_PNG)
            document, warnings = renderer.render(
                "![diagram](https://example.invalid/x.png)"
            )

        mock_urlopen.assert_called_once()
        self.assertEqual(warnings, [])

    def test_local_image_is_unaffected_by_allow_remote_images_false(self):
        renderer = GfmDocxRenderer("Arial", Pt(12), allow_remote_images=False)

        with tempfile.TemporaryDirectory() as directory:
            (Path(directory) / "diagram.png").write_bytes(_MINIMAL_PNG)
            with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
                document, warnings = renderer.render(
                    "![diagram](diagram.png)",
                    source_path=Path(directory) / "source.md",
                )

        mock_urlopen.assert_not_called()
        self.assertEqual(warnings, [])

    def test_is_remote_target_covers_http_unc_and_protocol_relative(self):
        self.assertTrue(_is_remote_target("http://example.com/x.png"))
        self.assertTrue(_is_remote_target("HTTPS://EXAMPLE.COM/X.PNG"))
        self.assertTrue(_is_remote_target("//attacker.example.com/share/a.png"))
        self.assertTrue(_is_remote_target(r"\\attacker.example.com\share\a.png"))
        self.assertFalse(_is_remote_target("images/diagram.png"))
        self.assertFalse(_is_remote_target("/absolute/local/diagram.png"))

    def test_protocol_relative_target_not_fetched_when_disallowed(self):
        renderer = _renderer_ready_for_direct_image_append(
            GfmDocxRenderer("Arial", Pt(12), allow_remote_images=False)
        )
        target = "//attacker.example.com/share/a.png"

        with patch("mdtoword.gfm_renderer.Path") as mock_path:
            renderer._append_image(_FakeImageToken(target), None)

        mock_path.assert_not_called()
        self.assertEqual(
            renderer.warnings,
            [
                f"Remote image not fetched: {target} "
                "(network access is disabled; pass fetch_remote_images=true to allow it)"
            ],
        )
        self.assertIn("[diagram]", renderer._paragraph.text)

    def test_unc_target_not_fetched_when_disallowed(self):
        renderer = _renderer_ready_for_direct_image_append(
            GfmDocxRenderer("Arial", Pt(12), allow_remote_images=False)
        )
        target = r"\\attacker.example.com\share\a.png"

        with patch("mdtoword.gfm_renderer.Path") as mock_path:
            renderer._append_image(_FakeImageToken(target), None)

        mock_path.assert_not_called()
        self.assertEqual(
            renderer.warnings,
            [
                f"Remote image not fetched: {target} "
                "(network access is disabled; pass fetch_remote_images=true to allow it)"
            ],
        )
        self.assertIn("[diagram]", renderer._paragraph.text)

    def test_unc_target_still_uses_filesystem_when_remote_images_allowed(self):
        # allow_remote_images defaults to True -- a UNC path must keep taking
        # the plain filesystem branch exactly as it did before this gate
        # existed, so a legitimate network share on Windows is unaffected.
        renderer = _renderer_ready_for_direct_image_append(GfmDocxRenderer("Arial", Pt(12)))
        target = r"\\attacker.example.com\share\a.png"

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            renderer._append_image(_FakeImageToken(target), None)

        mock_urlopen.assert_not_called()
        self.assertEqual(renderer.warnings, [f"Image not found: {target}"])
        self.assertIn("[diagram]", renderer._paragraph.text)

    def test_remote_image_over_size_cap_falls_back_with_warning(self):
        renderer = GfmDocxRenderer("Arial", Pt(12))
        oversized = b"x" * (_MAX_REMOTE_IMAGE_BYTES + 1)
        response_context = _urlopen_response(oversized)

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            mock_urlopen.return_value = response_context
            document, warnings = renderer.render(
                "![diagram](https://example.invalid/x.png)"
            )

        mock_urlopen.assert_called_once()
        # The mock's read() ignores what it is called with (it always
        # returns the full ``oversized`` buffer regardless), so the only way
        # to pin the actual memory bound -- reading at most one byte past the
        # cap, never the whole response -- is to assert the call argument
        # itself. Reverting to an unbounded ``response.read()`` would leave
        # every other assertion here green.
        response_context.__enter__.return_value.read.assert_called_once_with(
            _MAX_REMOTE_IMAGE_BYTES + 1
        )
        self.assertEqual(len(warnings), 1)
        self.assertIn("too large", warnings[0])
        self.assertIn(str(_MAX_REMOTE_IMAGE_BYTES), warnings[0])
        self.assertIn("[diagram]", document.paragraphs[0].text)

    def test_renders_footnotes_as_a_final_section(self):
        document, warnings = self.renderer.render(
            "Text with a footnote.[^1]\n\n[^1]: Footnote text."
        )

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Footnotes", text)
        self.assertIn("Footnote text.", text)
        self.assertEqual(warnings, [])

    def test_footnotes_heading_defaults_to_english(self):
        document, _ = self.renderer.render(
            "Text with a footnote.[^1]\n\n[^1]: Footnote text."
        )

        headings = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
        self.assertEqual(headings, ["Footnotes"])

    def test_footnotes_heading_can_be_localized(self):
        renderer = GfmDocxRenderer("Arial", Pt(12), footnotes_heading="Сноски")
        document, _ = renderer.render(
            "Text with a footnote.[^1]\n\n[^1]: Footnote text."
        )

        headings = [p.text for p in document.paragraphs if p.style.name == "Heading 2"]
        self.assertEqual(headings, ["Сноски"])

    def test_headings_are_black(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "# Заголовок\n\n## Второй\n\n### Третий\n"
        )
        for level in (1, 2, 3):
            style = document.styles[f"Heading {level}"]
            self.assertEqual(style.font.color.rgb, RGBColor(0, 0, 0))

    @staticmethod
    def _effective_size(paragraph):
        """Resolve rendered font size: run size wins, else the style's."""
        run = paragraph.runs[0]
        return run.font.size if run.font.size is not None else paragraph.style.font.size

    @staticmethod
    def _effective_bold(paragraph):
        run = paragraph.runs[0]
        return run.bold if run.bold is not None else paragraph.style.font.bold

    @staticmethod
    def _effective_italic(paragraph):
        run = paragraph.runs[0]
        return run.italic if run.italic is not None else paragraph.style.font.italic

    @staticmethod
    def _effective_font_name(paragraph):
        run = paragraph.runs[0]
        return run.font.name if run.font.name is not None else paragraph.style.font.name

    def test_heading_levels_get_a_real_size_hierarchy(self):
        document, _ = GfmDocxRenderer("Arial", Pt(12)).render(
            "# H1\n\n"
            "## H2\n\n"
            "### H3\n\n"
            "#### H4\n\n"
            "##### H5\n\n"
            "###### H6\n\n"
            "Body paragraph.\n"
        )

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "headings.docx"
            document.save(str(path))
            reopened = Document(str(path))

        headings = {
            paragraph.style.name: paragraph
            for paragraph in reopened.paragraphs
            if paragraph.style.name.startswith("Heading")
        }
        expected_pt = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
        for level, expected in expected_pt.items():
            paragraph = headings[f"Heading {level}"]
            self.assertEqual(
                self._effective_size(paragraph), Pt(expected),
                f"Heading {level} should render at {expected}pt",
            )
            self.assertTrue(
                self._effective_bold(paragraph), f"Heading {level} should be bold"
            )
            self.assertEqual(
                self._effective_font_name(paragraph), "Arial",
                f"Heading {level} should keep the constructor's font",
            )

        self.assertTrue(
            self._effective_italic(headings["Heading 6"]),
            "Heading 6 should be italic to stay distinguishable from Heading 5",
        )

        for level in range(1, 7):
            style = reopened.styles[f"Heading {level}"]
            self.assertEqual(style.font.color.rgb, RGBColor(0, 0, 0))

        body_paragraphs = [p for p in reopened.paragraphs if p.text == "Body paragraph."]
        self.assertEqual(len(body_paragraphs), 1)
        self.assertEqual(self._effective_size(body_paragraphs[0]), Pt(12))

    def test_heading_scale_is_relative_to_the_base_font_size(self):
        document, _ = GfmDocxRenderer("Arial", Pt(14)).render("# H1\n")

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scaled-heading.docx"
            document.save(str(path))
            reopened = Document(str(path))

        self.assertEqual(self._effective_size(reopened.paragraphs[0]), Pt(21))

    def test_hyperlink_is_black_and_underlined(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "[сайт](https://example.com)\n"
        )
        xml = document.paragraphs[0]._p.xml
        self.assertIn('w:val="000000"', xml)
        self.assertNotIn("0563C1", xml)
        self.assertIn("w:u", xml)

    def test_body_paragraphs_are_justified(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Обычный абзац текста.\n\n- пункт списка\n\n> цитата\n"
        )
        justified = [
            p for p in document.paragraphs
            if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        ]
        self.assertEqual(len(justified), 3)

    def test_headings_and_code_are_not_justified(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "# Заголовок\n\n```python\nx = 1\n```\n"
        )
        for paragraph in document.paragraphs:
            self.assertNotEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_quote_style_color_is_explicit_black_not_theme_color(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "> цитата\n"
        )
        style = document.styles["Quote"]
        self.assertEqual(style.font.color.rgb, RGBColor(0, 0, 0))
        self.assertNotIn("themeColor", style.element.xml)

    def test_table_has_explicit_borders(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "| a | b |\n|---|---|\n| 1 | 2 |\n"
        )
        table_xml = document.tables[0]._tbl.tblPr.xml
        self.assertIn("tblBorders", table_xml)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            self.assertIn(f"w:{edge}", table_xml)
        tbl_pr = document.tables[0]._tbl.tblPr
        child_names = [child.tag.rsplit("}", 1)[-1] for child in tbl_pr]
        self.assertIn("tblBorders", child_names)
        self.assertIn("tblLook", child_names)
        self.assertLess(
            child_names.index("tblBorders"),
            child_names.index("tblLook"),
            "w:tblBorders must precede w:tblLook in the tblPr child sequence "
            "per the OOXML CT_TblPrBase schema order",
        )

    def test_table_respects_markdown_column_alignment(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "| left | right | center |\n|:---|---:|:---:|\n| a | b | c |\n"
        )
        body = document.tables[0].rows[1]
        self.assertEqual(body.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(body.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(body.cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_table_cell_math_is_kept_verbatim_with_a_warning(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "| formula | plain |\n|---|---|\n| $x^2$ | c |\n"
        )
        cell = document.tables[0].cell(1, 0)
        self.assertEqual(cell.text, "$x^2$")
        self.assertEqual(_equations(cell.paragraphs[0]), [])
        self.assertEqual(len(warnings), 1)
        self.assertIn("table cell", warnings[0])

    def test_escaped_inline_math_becomes_equations_not_text(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"Формула $a\,b$ и $a*b*c$ и $50\%$ в строке."
        )
        paragraph = document.paragraphs[0]
        self.assertEqual(len(_equations(paragraph)), 3)
        self.assertEqual(paragraph.text, "Формула  и  и  в строке.")
        self.assertEqual(warnings, [])

    def test_dollar_amounts_in_prose_survive_verbatim(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "This costs $5 and that costs $10, total $15."
        )
        text = document.paragraphs[0].text
        self.assertEqual(
            text, "This costs $5 and that costs $10, total $15."
        )
        self.assertEqual(warnings, [])

    def test_display_math_becomes_a_structured_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\n\\int_0^\\infty e^{-x^2}\\,dx = \\frac{\\sqrt{\\pi}}{2}\n$$\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        xml = with_equations[0]._p.xml
        self.assertIn("<m:nary>", xml)  # the integral
        self.assertIn("<m:f>", xml)  # the fraction
        self.assertIn("<m:rad>", xml)  # the square root
        self.assertEqual(warnings, [])

    def test_equation_label_is_not_dropped(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\nx = 1\n$$ (1)\n"
        )
        labelled = [p for p in document.paragraphs if "(1)" in p.text]
        self.assertEqual(len(labelled), 1, "equation label was dropped")
        self.assertEqual(
            len(_equations(labelled[0])),
            1,
            "label paragraph lost its equation",
        )
        self.assertEqual(warnings, [])

    def test_amsmath_align_becomes_one_equation_aligned_on_its_ampersands(self):
        r"""``align`` converts whole: one centred Word equation holding an
        equation array, with an alignment point on the run each ``&``
        precedes -- so the two ``=`` line up the way LaTeX draws them,
        instead of the environment being cut into one paragraph per line
        with the alignment discarded."""
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{align}\n"
            "a &= b + c \\\\\n"
            "d &= e + f\n"
            "\\end{align}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        paragraph = with_equations[0]
        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(len(_equations(paragraph)), 1)
        lines = _equations(paragraph)[0].findall(f"{_MATH_NS}eqArr/{_MATH_NS}e")
        self.assertEqual(len(lines), 2)
        for line in lines:
            aligned = [
                "".join(t.text or "" for t in run.iter(f"{_MATH_NS}t"))
                for run in line.iter(f"{_MATH_NS}r")
                if run.find(f"{_MATH_NS}rPr/{_MATH_NS}aln") is not None
            ]
            self.assertEqual(aligned, ["="])
        self.assertEqual(warnings, [])

    def test_amsmath_matrix_environment_is_a_single_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{pmatrix}\n"
            "a & b \\\\\n"
            "c & d\n"
            "\\end{pmatrix}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        self.assertEqual(len(_equations(with_equations[0])), 1)
        self.assertIn("<m:m>", with_equations[0]._p.xml)
        self.assertEqual(warnings, [])

    def test_amsmath_equation_keeps_a_leading_brace_group(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{equation}{a} + b = c\\end{equation}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        equations = _equations(with_equations[0])
        self.assertEqual(len(equations), 1)
        text = _equation_text(equations[0])
        for expected in ("a", "+", "b", "=", "c"):
            self.assertIn(expected, text)
        self.assertEqual(warnings, [])

    def test_amsmath_gather_keeps_a_leading_brace_group(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{gather}{a} + b\\end{gather}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        text = _equation_text(_equations(with_equations[0])[0])
        for expected in ("a", "+", "b"):
            self.assertIn(expected, text)
        self.assertEqual(warnings, [])

    def test_amsmath_align_keeps_a_leading_brace_group(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{align}{a} &= b\\end{align}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        text = _equation_text(_equations(with_equations[0])[0])
        for expected in ("a", "=", "b"):
            self.assertIn(expected, text)
        self.assertEqual(warnings, [])

    def test_amsmath_alignat_consumes_its_column_argument(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{alignat}{2}\n"
            "a &= b \\\\\n"
            "c &= d\n"
            "\\end{alignat}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        combined_text = "".join(
            _equation_text(equation)
            for paragraph in with_equations
            for equation in _equations(paragraph)
        )
        self.assertNotIn("2", combined_text)
        for expected in ("a", "b", "c", "d"):
            self.assertIn(expected, combined_text)
        self.assertEqual(warnings, [])

    def test_amsmath_flalign_converts_to_equations(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{flalign}\n"
            "a &= b \\\\\n"
            "c &= d\n"
            "\\end{flalign}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertGreaterEqual(len(with_equations), 1)
        self.assertEqual(warnings, [])

    def test_amsmath_gather_converts_to_equations(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{gather}\n"
            "a + b \\\\\n"
            "c + d\n"
            "\\end{gather}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertGreaterEqual(len(with_equations), 1)
        self.assertEqual(warnings, [])

    def test_amsmath_multline_converts_to_equations(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{multline}\n"
            "a + b \\\\\n"
            "c + d\n"
            "\\end{multline}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertGreaterEqual(len(with_equations), 1)
        self.assertEqual(warnings, [])

    def test_amsmath_eqnarray_converts_to_equations(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{eqnarray}\n"
            "a &= b \\\\\n"
            "c &= d\n"
            "\\end{eqnarray}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertGreaterEqual(len(with_equations), 1)
        self.assertEqual(warnings, [])

    def test_amsmath_equation_with_nested_matrix_is_one_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{equation}\n"
            "A = \\begin{pmatrix} a & b \\\\ c & d \\end{pmatrix}\n"
            "\\end{equation}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        self.assertEqual(len(_equations(with_equations[0])), 1)
        self.assertIn("<m:m>", with_equations[0]._p.xml)
        self.assertEqual(warnings, [])

    def test_display_math_line_break_becomes_one_stacked_equation(self):
        r"""``\\`` no longer needs an amsmath environment around it: a
        display formula that uses it stays a single Word equation whose
        lines stack, rather than being kept as text with a warning."""
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\na = b \\\\ c = d\n$$\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        equation = _equations(with_equations[0])[0]
        lines = equation.findall(f"{_MATH_NS}eqArr/{_MATH_NS}e")
        self.assertEqual(len(lines), 2)
        self.assertEqual(warnings, [])

    def test_inline_math_line_break_converts_without_a_warning(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Обе строки: $a \\\\ b$ здесь.\n"
        )
        self.assertEqual(len(_equations(document.paragraphs[0])), 1)
        self.assertEqual(warnings, [])

    def test_amsmath_gather_becomes_one_equation_with_stacked_lines(self):
        r"""``gather`` has no ``&`` alignment, so the whole body now
        converts in one piece -- a single Word equation whose ``\\`` lines
        stack -- instead of falling back to one paragraph per line."""
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{gather}\n"
            "a + b \\\\\n"
            "c + d\n"
            "\\end{gather}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        lines = _equations(with_equations[0])[0].findall(
            f"{_MATH_NS}eqArr/{_MATH_NS}e"
        )
        self.assertEqual(
            ["".join(t.text or "" for t in line.iter(f"{_MATH_NS}t"))
             for line in lines],
            ["a+b", "c+d"],
        )
        self.assertEqual(warnings, [])

    def test_single_line_align_still_goes_through_the_split_fallback(self):
        r"""A one-line ``align`` has no second line to align against, so
        ``latex_omml`` refuses its ``&`` as a probable unescaped ampersand
        and the whole-body attempt fails. That is what keeps the ``\\``
        splitting path alive: it strips the ``&`` and still produces one
        equation rather than dropping the environment to verbatim text."""
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{align}a &= b\\end{align}\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        text = _equation_text(_equations(with_equations[0])[0])
        for expected in ("a", "=", "b"):
            self.assertIn(expected, text)
        self.assertEqual(warnings, [])

    def test_display_array_environment_becomes_an_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\n\\begin{array}{lr} a & b \\\\ c & d \\end{array}\n$$\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        self.assertIn("<m:m>", with_equations[0]._p.xml)
        self.assertIn('<m:mcJc m:val="left"/>', with_equations[0]._p.xml)
        self.assertEqual(warnings, [])

    def test_array_with_a_vertical_rule_is_kept_as_text_and_warns(self):
        r"""The one array shape OMML cannot express must still land in the
        document character for character, with a warning naming why."""
        source = r"\begin{array}{c|c} a & b \end{array}"
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            f"$$\n{source}\n$$\n"
        )
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn(source, text)
        self.assertEqual(
            [eq for p in document.paragraphs for eq in _equations(p)], [])
        self.assertEqual(len(warnings), 1)
        self.assertIn("Formula kept as text", warnings[0])
        self.assertIn("|", warnings[0])

    def test_display_choose_and_substack_convert_without_warning(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\\sum_{\\substack{i \\\\ j}} {n \\choose k}$$\n"
        )
        with_equations = [p for p in document.paragraphs if _equations(p)]
        self.assertEqual(len(with_equations), 1)
        xml = with_equations[0]._p.xml
        self.assertIn("<m:nary>", xml)
        self.assertIn('<m:type m:val="noBar"/>', xml)
        self.assertEqual(warnings, [])

    def test_unsupported_display_environment_is_kept_byte_for_byte(self):
        source = (
            "\\begin{align}\n"
            "a &= \\qedsymbol{b} \\\\\n"
            "c &= d\n"
            "\\end{align}"
        )
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            source + "\n"
        )
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertEqual(text, source)
        self.assertEqual(
            [eq for p in document.paragraphs for eq in _equations(p)],
            [],
            "a partially converted environment must not leak equations",
        )
        self.assertEqual(len(warnings), 1)
        self.assertIn("qedsymbol", warnings[0])

    def test_whitespace_only_inline_math_keeps_its_literal_dollars(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "x $ $ y\n"
        )
        self.assertEqual(document.paragraphs[0].text, "x $ $ y")
        self.assertEqual(_equations(document.paragraphs[0]), [])
        self.assertEqual(warnings, [])

    def test_empty_display_math_does_not_crash_or_warn(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "before\n\n$$\n$$\n\nafter\n"
        )
        text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("before", text)
        self.assertIn("after", text)
        self.assertEqual(warnings, [])

    def test_blank_line_does_not_widen_display_math(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Первый абзац.\n"
            "\n"
            "$$\n"
            "\n"
            "Второй абзац, ничего общего с формулой.\n"
            "\n"
            "$$\n"
            "\n"
            "Третий абзац.\n"
        )
        middle_text = "Второй абзац, ничего общего с формулой."
        matches = [p for p in document.paragraphs if middle_text in p.text]
        self.assertTrue(matches, "middle paragraph text did not survive verbatim")
        middle_paragraph = matches[0]
        self.assertTrue(middle_paragraph.runs, "middle paragraph has no runs")
        for run in middle_paragraph.runs:
            self.assertNotEqual(run.font.name, "Courier New")

    def test_stray_dollar_pair_with_cyrillic_content_warns(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Переменные $HOME и $PATH в шелле."
        )
        self.assertEqual(len(warnings), 1)
        self.assertIn(r"\$", warnings[0])
        self.assertIn("HOME", warnings[0])
        self.assertIn("$HOME и $", document.paragraphs[0].text)

    def test_genuine_formula_does_not_warn(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$E = mc^2$"
        )
        self.assertEqual(warnings, [])

    def test_text_wrapped_cyrillic_converts_to_a_real_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"$\text{путь}$"
        )
        self.assertEqual(len(_equations(document.paragraphs[0])), 1)
        self.assertEqual(warnings, [])

    def test_text_wrapped_cyrillic_inside_a_fraction_converts_with_no_warning(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"$v = \frac{\text{путь}}{\text{время}}$"
        )
        paragraph = document.paragraphs[0]
        equations = _equations(paragraph)
        self.assertEqual(len(equations), 1)
        self.assertIn("<m:f>", paragraph._p.xml)
        self.assertEqual(warnings, [])

    def test_shell_variable_dollar_pair_still_warns(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$HOME и $PATH"
        )
        self.assertEqual(len(warnings), 1)
        self.assertIn("Cyrillic", warnings[0])
        self.assertEqual(_equations(document.paragraphs[0]), [])

    def test_english_prose_dollar_pair_survives_verbatim(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Set $PATH and $HOME before running.\n"
        )
        self.assertEqual(
            document.paragraphs[0].text, "Set $PATH and $HOME before running."
        )
        self.assertEqual(_equations(document.paragraphs[0]), [])
        self.assertEqual(len(warnings), 1)

    def test_english_price_range_dollar_pair_survives_verbatim(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Prices ranged from $low to $high in Q3.\n"
        )
        self.assertEqual(
            document.paragraphs[0].text,
            "Prices ranged from $low to $high in Q3.",
        )
        self.assertEqual(_equations(document.paragraphs[0]), [])
        self.assertEqual(len(warnings), 1)

    def test_genuine_formula_regression_battery_still_converts(self):
        formulas = [
            r"$a+b$",
            r"$x^2$",
            r"$E = mc^2$",
            r"$\alpha$",
            r"$\frac{a}{b}$",
            r"$n$",
            r"$xy$",
        ]
        for formula in formulas:
            with self.subTest(formula=formula):
                document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
                    formula + "\n"
                )
                self.assertEqual(
                    warnings, [], f"{formula} should convert without a warning"
                )
                self.assertEqual(
                    len(_equations(document.paragraphs[0])), 1,
                    f"{formula} should become a real equation",
                )

    def test_bare_cyrillic_alongside_text_command_still_warns(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"$\text{путь} + скорость$"
        )
        self.assertEqual(len(warnings), 1)
        self.assertIn("Cyrillic", warnings[0])
        self.assertEqual(_equations(document.paragraphs[0]), [])

    def test_inline_math_becomes_a_real_word_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"Энергия $E = mc^2$ покоя."
        )
        xml = document.paragraphs[0]._p.xml
        self.assertIn("oMath", xml)
        self.assertIn("sSup", xml)
        self.assertEqual(warnings, [])

    def test_display_math_becomes_a_centred_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "$$\\frac{a}{b}$$\n"
        )
        xml = "".join(p._p.xml for p in document.paragraphs)
        self.assertIn("oMath", xml)
        self.assertIn("<m:f>", xml)
        self.assertEqual(warnings, [])

    def test_unsupported_formula_falls_back_to_verbatim_text_and_warns(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            r"Формула $\qedsymbol{x}$ здесь."
        )
        self.assertIn(r"$\qedsymbol{x}$", document.paragraphs[0].text)
        self.assertEqual(len(warnings), 1)
        self.assertIn("qedsymbol", warnings[0])

    def test_amsmath_equation_becomes_an_equation(self):
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "\\begin{equation}\nF = G\\frac{m_1 m_2}{r^2}\n\\end{equation}\n"
        )
        xml = "".join(p._p.xml for p in document.paragraphs)
        self.assertIn("oMath", xml)
        self.assertEqual(warnings, [])

    def test_equations_survive_a_save_and_reopen_round_trip(self):
        unsupported = r"\qedsymbol{x}"
        document, warnings = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "Inline $E = mc^2$ here.\n\n"
            "$$\\frac{a}{b}$$\n\n"
            "\\begin{equation}\n"
            "F = G\\frac{m_1 m_2}{r^2}\n"
            "\\end{equation}\n\n"
            "\\begin{align}\n"
            "a &= b + c \\\\\n"
            "d &= e + f\n"
            "\\end{align}\n\n"
            f"Broken ${unsupported}$ formula.\n"
        )

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "round-trip.docx"
            document.save(str(path))
            reopened = Document(str(path))

        equations = [eq for p in reopened.paragraphs for eq in _equations(p)]
        # inline + display + equation + the align environment, which is one
        # equation holding both of its lines rather than one per line.
        self.assertEqual(len(equations), 4)

        self.assertEqual(len(warnings), 1)
        self.assertIn("qedsymbol", warnings[0])

        text = "\n".join(p.text for p in reopened.paragraphs)
        self.assertIn(f"Broken ${unsupported}$ formula.", text)

    def test_footnote_paragraphs_are_justified_like_body_lists(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "1. пункт списка\n\n"
            "Текст со сноской[^1]\n\n"
            "[^1]: Содержимое сноски.\n"
        )
        list_number_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "List Number"
        ]
        self.assertEqual(len(list_number_paragraphs), 2)
        for paragraph in list_number_paragraphs:
            self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_every_applied_style_uses_the_chosen_font_with_no_theme_override(self):
        document, _ = GfmDocxRenderer("Times New Roman", Pt(12)).render(
            "# H1\n\n"
            "## H2\n\n"
            "### H3\n\n"
            "#### H4\n\n"
            "##### H5\n\n"
            "###### H6\n\n"
            "Body paragraph.\n\n"
            "> quoted text\n\n"
            "- bullet item\n\n"
            "1. numbered item\n"
        )

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "theme-fonts.docx"
            document.save(str(path))
            reopened = Document(str(path))

        for name in _STYLES_THE_RENDERER_APPLIES:
            style = reopened.styles[name]
            attrs = _rfonts_attrs(style)
            self.assertIsNotNone(attrs, f'"{name}" style has no w:rFonts element')
            for attribute_name in attrs:
                self.assertFalse(
                    attribute_name.lower().endswith("theme"),
                    f'"{name}" w:rFonts still carries a theme attribute: '
                    f"{attribute_name}={attrs[attribute_name]!r}",
                )
            self.assertEqual(
                attrs.get("ascii"), "Times New Roman",
                f'"{name}" w:rFonts/@w:ascii should be the requested font',
            )
            self.assertEqual(
                attrs.get("hAnsi"), "Times New Roman",
                f'"{name}" w:rFonts/@w:hAnsi should be the requested font',
            )

    def test_heading_font_comes_from_the_constructor_not_a_hardcoded_theme_font(self):
        document, _ = GfmDocxRenderer("Georgia", Pt(12)).render(
            "# H1\n\n"
            "## H2\n\n"
            "### H3\n\n"
            "#### H4\n\n"
            "##### H5\n\n"
            "###### H6\n"
        )

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "georgia-headings.docx"
            document.save(str(path))
            reopened = Document(str(path))

        for level in range(1, 7):
            style = reopened.styles[f"Heading {level}"]
            attrs = _rfonts_attrs(style)
            self.assertIsNotNone(attrs, f"Heading {level} style has no w:rFonts element")
            self.assertEqual(
                attrs.get("ascii"), "Georgia",
                f"Heading {level} should take its font from the constructor",
            )
            self.assertEqual(
                attrs.get("hAnsi"), "Georgia",
                f"Heading {level} should take its font from the constructor",
            )

    def test_absolute_target_outside_root_is_refused_without_probing_filesystem(self):
        # The oracle-closing property: containment must be decided before
        # is_file() is ever called on the candidate path, since the answer
        # to "does this exist" is exactly what an agent could extract one
        # bit at a time by naming arbitrary local files in the Markdown.
        with tempfile.TemporaryDirectory() as root_dir, \
                tempfile.TemporaryDirectory() as outside_dir:
            root = Path(root_dir)
            outside_target = str(Path(outside_dir) / "secret.png")
            renderer = _renderer_ready_for_direct_image_append(
                GfmDocxRenderer("Arial", Pt(12), image_roots=[root])
            )

            with patch.object(Path, "is_file", autospec=True) as mock_is_file:
                renderer._append_image(_FakeImageToken(outside_target), None)

        mock_is_file.assert_not_called()
        self.assertEqual(
            renderer.warnings,
            [
                f"Image outside the allowed root: {outside_target} "
                "(pass image_root=... to widen it)"
            ],
        )
        self.assertIn("[diagram]", renderer._paragraph.text)

    def test_refusal_message_is_identical_whether_or_not_the_file_exists(self):
        # Same target, checked both before and after the file is created --
        # if the message depended on existence, this would catch it.
        with tempfile.TemporaryDirectory() as root_dir, \
                tempfile.TemporaryDirectory() as outside_dir:
            root = Path(root_dir)
            outside_target = str(Path(outside_dir) / "secret.png")
            renderer = _renderer_ready_for_direct_image_append(
                GfmDocxRenderer("Arial", Pt(12), image_roots=[root])
            )

            renderer._append_image(_FakeImageToken(outside_target), None)
            missing_warning = renderer.warnings[-1]

            Path(outside_target).write_bytes(_MINIMAL_PNG)
            renderer.warnings = []
            renderer._append_image(_FakeImageToken(outside_target), None)
            existing_warning = renderer.warnings[-1]

        self.assertEqual(missing_warning, existing_warning)
        self.assertEqual(
            missing_warning,
            f"Image outside the allowed root: {outside_target} "
            "(pass image_root=... to widen it)",
        )

    def test_relative_traversal_that_stays_inside_root_still_embeds(self):
        with tempfile.TemporaryDirectory() as root_dir:
            root = Path(root_dir)
            (root / "images").mkdir()
            (root / "images" / "logo.png").write_bytes(_MINIMAL_PNG)
            (root / "guide").mkdir()
            source_path = root / "guide" / "source.md"

            renderer = _renderer_ready_for_direct_image_append(
                GfmDocxRenderer("Arial", Pt(12), image_roots=[root])
            )
            renderer._append_image(_FakeImageToken("../images/logo.png"), source_path)

        self.assertEqual(renderer.warnings, [])
        self.assertEqual(len(renderer.document.inline_shapes), 1)

    def test_traversal_that_escapes_root_is_refused(self):
        with tempfile.TemporaryDirectory() as root_dir:
            root = Path(root_dir)
            (root / "guide").mkdir()
            source_path = root / "guide" / "source.md"
            target = "../../../etc/hosts"

            renderer = _renderer_ready_for_direct_image_append(
                GfmDocxRenderer("Arial", Pt(12), image_roots=[root])
            )
            renderer._append_image(_FakeImageToken(target), source_path)

        self.assertEqual(
            renderer.warnings,
            [f"Image outside the allowed root: {target} (pass image_root=... to widen it)"],
        )

    def test_symlink_inside_root_pointing_outside_is_refused(self):
        # A prefix check on the unresolved path would be fooled by this --
        # the string "root/link.png" starts with "root", even though the
        # symlink it names resolves somewhere else entirely.
        with tempfile.TemporaryDirectory() as root_dir, \
                tempfile.TemporaryDirectory() as outside_dir:
            root = Path(root_dir)
            outside_target = Path(outside_dir) / "secret.png"
            outside_target.write_bytes(_MINIMAL_PNG)
            link = root / "link.png"
            link.symlink_to(outside_target)

            renderer = _renderer_ready_for_direct_image_append(
                GfmDocxRenderer("Arial", Pt(12), image_roots=[root])
            )
            renderer._append_image(_FakeImageToken("link.png"), root / "doc.md")

        self.assertEqual(
            renderer.warnings,
            ["Image outside the allowed root: link.png (pass image_root=... to widen it)"],
        )

    def test_image_roots_none_is_unrestricted(self):
        # Pins the GUI-preservation guarantee: app.py never passes
        # image_roots, so the default must keep behaving exactly as it did
        # before this restriction existed.
        with tempfile.TemporaryDirectory() as directory:
            image_path = Path(directory) / "diagram.png"
            image_path.write_bytes(_MINIMAL_PNG)

            renderer = _renderer_ready_for_direct_image_append(GfmDocxRenderer("Arial", Pt(12)))
            renderer._append_image(_FakeImageToken(str(image_path)), None)

        self.assertEqual(renderer.warnings, [])
        self.assertEqual(len(renderer.document.inline_shapes), 1)


if __name__ == "__main__":
    unittest.main()
