"""Тесты MCP-сервера через in-memory клиент из SDK.

Клиент и сервер соединяются напрямую в одном процессе, без подпроцесса
и без stdio, — проверяется реальный путь вызова инструмента вместе со
схемами и валидацией аргументов.
"""

from pathlib import Path
import os
import tempfile
import unittest
from unittest.mock import MagicMock, patch

from docx import Document

try:
    from mcp.shared.memory import (
        create_connected_server_and_client_session as client_session,
    )

    from mdtoword.mcp_server import mcp as server
except ImportError:  # pragma: no cover
    client_session = None
    server = None


# The smallest well-formed PNG, used so a mocked remote fetch has something
# valid to embed rather than tripping the renderer's own error handling.
_MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\x05-\xb4\x00\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _urlopen_response(data: bytes) -> MagicMock:
    """Mock of ``urlopen(...)``'s return value, usable as a context manager."""
    response = MagicMock()
    response.read.return_value = data
    context = MagicMock()
    context.__enter__.return_value = response
    context.__exit__.return_value = False
    return context


def setUpModule() -> None:
    """Пропустить набор целиком, если SDK mcp не установлен.

    Пропуск объявлен здесь, а не прямо на уровне модуля: ``python -m unittest``
    не перехватывает ``SkipTest``, выброшенный во время импорта, и обрывает
    весь прогон, а README документирует именно эту команду. ``setUpModule``
    же корректно понимают оба раннера.
    """
    if server is None:  # pragma: no cover
        raise unittest.SkipTest("SDK mcp не установлен; см. requirements-mcp.txt")


class McpServerTestCase(unittest.IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        self._tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmpdir.cleanup)
        self.root = Path(self._tmpdir.name)

    async def call(self, tool: str, arguments: dict):
        async with client_session(server._mcp_server) as client:
            return await client.call_tool(tool, arguments)


class ToolRegistrationTests(McpServerTestCase):
    async def test_both_conversion_tools_are_advertised_with_descriptions(self) -> None:
        async with client_session(server._mcp_server) as client:
            listed = await client.list_tools()

        tools = {tool.name: tool for tool in listed.tools}
        self.assertIn("markdown_to_word", tools)
        self.assertIn("word_to_markdown", tools)
        for tool in tools.values():
            self.assertTrue(tool.description)

    async def test_the_lossy_direction_says_so_in_its_description(self) -> None:
        async with client_session(server._mcp_server) as client:
            listed = await client.list_tools()

        description = next(t.description for t in listed.tools if t.name == "word_to_markdown")
        self.assertIn("lossy", description.lower())


class MarkdownToWordTests(McpServerTestCase):
    async def test_a_directory_is_converted_recursively_with_one_output_each(self) -> None:
        nested = self.root / "nested"
        nested.mkdir()
        (self.root / "first.md").write_text("# Первый", encoding="utf-8")
        (nested / "second.markdown").write_text("# Второй", encoding="utf-8")
        (self.root / "ignored.txt").write_text("не markdown", encoding="utf-8")

        result = await self.call("markdown_to_word", {"inputs": [str(self.root)]})

        self.assertFalse(result.isError)
        report = result.structuredContent
        self.assertEqual(report["sources_found"], 2)
        self.assertEqual(len(report["converted"]), 2)
        self.assertEqual(report["failed"], [])
        for entry in report["converted"]:
            self.assertTrue(Path(entry["output"]).is_file())

    async def test_output_dir_is_created_and_used(self) -> None:
        (self.root / "doc.md").write_text("# Заголовок", encoding="utf-8")
        destination = self.root / "out" / "deep"

        result = await self.call(
            "markdown_to_word",
            {"inputs": [str(self.root / "doc.md")], "output_dir": str(destination)},
        )

        report = result.structuredContent
        # .resolve() с обеих сторон: на macOS временный каталог лежит под /var,
        # который является симлинком на /private/var, а _prepare_output_dir
        # теперь резолвит output_dir, поэтому путь в отчёте — каноническая
        # форма destination.
        self.assertEqual(
            Path(report["converted"][0]["output"]).parent, destination.resolve()
        )
        self.assertTrue((destination / "doc.docx").is_file())

    async def test_relative_output_dir_is_resolved_to_an_absolute_path(self) -> None:
        self.addCleanup(os.chdir, os.getcwd())
        (self.root / "doc.md").write_text("# Заголовок", encoding="utf-8")
        os.chdir(self.root)

        result = await self.call(
            "markdown_to_word",
            {"inputs": ["doc.md"], "output_dir": "out"},
        )

        report = result.structuredContent
        output = Path(report["converted"][0]["output"])
        self.assertTrue(output.is_absolute())
        self.assertTrue(output.is_file())

    async def test_nonfatal_warnings_are_reported_per_file(self) -> None:
        (self.root / "doc.md").write_text("![diagram](missing.png)", encoding="utf-8")

        result = await self.call("markdown_to_word", {"inputs": [str(self.root)]})

        report = result.structuredContent
        self.assertEqual(
            report["converted"][0]["warnings"], ["Image not found: missing.png"]
        )

    async def test_font_arguments_reach_the_document(self) -> None:
        (self.root / "doc.md").write_text("Текст", encoding="utf-8")

        result = await self.call(
            "markdown_to_word",
            {"inputs": [str(self.root)], "font_name": "Georgia", "font_size": 14},
        )

        output = Path(result.structuredContent["converted"][0]["output"])
        document = Document(str(output))
        self.assertEqual(document.styles["Normal"].font.name, "Georgia")

    async def test_paths_matching_nothing_report_zero_sources_found(self) -> None:
        result = await self.call(
            "markdown_to_word", {"inputs": [str(self.root / "нет-такой-папки")]}
        )

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 0)
        self.assertEqual(report["converted"], [])
        self.assertEqual(report["failed"], [])

    async def test_output_dir_is_not_created_when_nothing_matches(self) -> None:
        destination = self.root / "out" / "deep"

        result = await self.call(
            "markdown_to_word",
            {
                "inputs": [str(self.root / "нет-такой-папки")],
                "output_dir": str(destination),
            },
        )

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 0)
        self.assertFalse(destination.exists())

    async def test_empty_inputs_is_an_error_not_an_empty_success(self) -> None:
        result = await self.call("markdown_to_word", {"inputs": []})

        self.assertTrue(result.isError)

    async def test_default_does_not_fetch_remote_images(self) -> None:
        (self.root / "doc.md").write_text(
            "![diagram](https://example.invalid/x.png)", encoding="utf-8"
        )

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            result = await self.call("markdown_to_word", {"inputs": [str(self.root)]})

        mock_urlopen.assert_not_called()
        warnings = result.structuredContent["converted"][0]["warnings"]
        self.assertEqual(len(warnings), 1)
        self.assertIn("fetch_remote_images", warnings[0])


class WordToMarkdownTests(McpServerTestCase):
    def write_docx(self, name: str) -> Path:
        path = self.root / name
        document = Document()
        document.add_heading("Раздел", level=1)
        document.save(str(path))
        return path

    async def test_documents_are_converted_to_markdown_files(self) -> None:
        self.write_docx("report.docx")

        result = await self.call("word_to_markdown", {"inputs": [str(self.root)]})

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 1)
        output = Path(report["converted"][0]["output"])
        self.assertEqual(output.suffix, ".md")
        self.assertIn("# Раздел", output.read_text(encoding="utf-8"))

    async def test_a_broken_file_fails_alone_without_stopping_the_batch(self) -> None:
        self.write_docx("good.docx")
        (self.root / "broken.docx").write_text("это не zip-контейнер", encoding="utf-8")

        result = await self.call("word_to_markdown", {"inputs": [str(self.root)]})

        report = result.structuredContent
        self.assertFalse(result.isError)
        self.assertEqual(report["sources_found"], 2)
        self.assertEqual(len(report["converted"]), 1)
        self.assertEqual(len(report["failed"]), 1)
        self.assertTrue(report["failed"][0]["source"].endswith("broken.docx"))
        self.assertTrue(report["failed"][0]["error"])

    async def test_output_dir_is_not_created_when_nothing_matches(self) -> None:
        destination = self.root / "out" / "deep"

        result = await self.call(
            "word_to_markdown",
            {
                "inputs": [str(self.root / "нет-такой-папки")],
                "output_dir": str(destination),
            },
        )

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 0)
        self.assertFalse(destination.exists())


class PreviewTests(McpServerTestCase):
    async def test_preview_reports_warnings_and_writes_no_files(self) -> None:
        (self.root / "doc.md").write_text("![diagram](missing.png)", encoding="utf-8")
        before = sorted(path.name for path in self.root.iterdir())

        result = await self.call("preview_markdown", {"inputs": [str(self.root)]})

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 1)
        self.assertEqual(
            report["previews"][0]["warnings"], ["Image not found: missing.png"]
        )
        self.assertEqual(sorted(path.name for path in self.root.iterdir()), before)

    async def test_preview_reports_unreadable_files_in_failed(self) -> None:
        (self.root / "good.md").write_text("# Заголовок", encoding="utf-8")
        broken = self.root / "broken.md"
        broken.write_bytes(b"\xff\xfe\x00 invalid utf-8")

        result = await self.call("preview_markdown", {"inputs": [str(self.root)]})

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 2)
        self.assertEqual(len(report["previews"]), 1)
        self.assertEqual(len(report["failed"]), 1)
        self.assertTrue(report["failed"][0]["source"].endswith("broken.md"))

    async def test_paths_matching_nothing_report_zero_sources_found(self) -> None:
        result = await self.call(
            "preview_markdown", {"inputs": [str(self.root / "нет-такой-папки")]}
        )

        report = result.structuredContent
        self.assertEqual(report["sources_found"], 0)
        self.assertEqual(report["previews"], [])
        self.assertEqual(report["failed"], [])

    async def test_default_does_not_fetch_remote_images(self) -> None:
        (self.root / "doc.md").write_text(
            "![diagram](https://example.invalid/x.png)", encoding="utf-8"
        )

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            result = await self.call("preview_markdown", {"inputs": [str(self.root)]})

        mock_urlopen.assert_not_called()
        warnings = result.structuredContent["previews"][0]["warnings"]
        self.assertEqual(len(warnings), 1)
        self.assertIn("fetch_remote_images", warnings[0])

    async def test_fetch_remote_images_true_fetches(self) -> None:
        (self.root / "doc.md").write_text(
            "![diagram](https://example.invalid/x.png)", encoding="utf-8"
        )

        with patch("mdtoword.gfm_renderer.urlopen") as mock_urlopen:
            mock_urlopen.return_value = _urlopen_response(_MINIMAL_PNG)
            result = await self.call(
                "preview_markdown",
                {"inputs": [str(self.root)], "fetch_remote_images": True},
            )

        mock_urlopen.assert_called_once()
        self.assertEqual(result.structuredContent["previews"][0]["warnings"], [])


class StdioProtocolTests(unittest.TestCase):
    def test_importing_the_server_writes_nothing_to_stdout(self) -> None:
        # stdout — это канал stdio-протокола: одна лишняя строка при импорте
        # рвёт JSON-RPC сессию, и клиент видит нечитаемую ошибку парсинга.
        import subprocess
        import sys

        repo_root = Path(__file__).resolve().parent.parent
        result = subprocess.run(
            [sys.executable, "-c", "import mdtoword.mcp_server"],
            capture_output=True,
            text=True,
            cwd=repo_root,
        )

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertEqual(result.stdout, "")


if __name__ == "__main__":
    unittest.main()
