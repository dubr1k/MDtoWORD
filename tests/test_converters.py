"""Тесты ядра конвертации.

Модуль намеренно не импортирует ``mdtoword.app``: ядро не должно зависеть
от PyQt6, и этот файл — исполняемое доказательство того, что не зависит.
"""

from pathlib import Path
import subprocess
import sys
import tempfile
import unittest

from docx import Document

from mdtoword.converters import (
    ConversionError,
    MarkdownToWordConverter,
    WordToMarkdownConverter,
)


class MarkdownToWordConverterTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmpdir.cleanup)
        self.root = Path(self._tmpdir.name)

    def test_convert_content_writes_the_document_and_returns_no_warnings(self) -> None:
        output = self.root / "out.docx"

        warnings = MarkdownToWordConverter().convert_content("# Заголовок", output)

        self.assertEqual(warnings, [])
        self.assertEqual(Document(str(output)).paragraphs[0].text, "Заголовок")

    def test_convert_file_returns_nonfatal_renderer_warnings_as_a_list(self) -> None:
        source = self.root / "source.md"
        source.write_text("![diagram](missing.png)", encoding="utf-8")
        output = self.root / "source.docx"

        warnings = MarkdownToWordConverter().convert_file(source, output)

        self.assertEqual(warnings, ["Image not found: missing.png"])
        self.assertTrue(output.is_file())

    def test_missing_source_raises_conversion_error(self) -> None:
        missing = self.root / "нет-такого.md"
        with self.assertRaises(OSError) as underlying:
            missing.read_text(encoding="utf-8")

        with self.assertRaises(ConversionError) as ctx:
            MarkdownToWordConverter().convert_file(missing, self.root / "out.docx")

        # Сообщение — это текст исходной ошибки без добавленного локализованного
        # префикса: язык подставляет потребитель (см. docstring модуля).
        self.assertEqual(str(ctx.exception), str(underlying.exception))

    def test_non_utf8_source_raises_conversion_error(self) -> None:
        source = self.root / "source.md"
        source.write_bytes(b"\xff\xfe\x00 invalid utf-8")

        with self.assertRaises(ConversionError):
            MarkdownToWordConverter().convert_file(source, self.root / "out.docx")

    def test_unwritable_output_raises_conversion_error(self) -> None:
        source = self.root / "source.md"
        source.write_text("# Заголовок", encoding="utf-8")
        # Каталог вместо файла: python-docx не может писать поверх директории.
        blocked = self.root / "blocked.docx"
        blocked.mkdir()

        with self.assertRaises(ConversionError):
            MarkdownToWordConverter().convert_file(source, blocked)

    def test_font_settings_reach_the_rendered_document(self) -> None:
        output = self.root / "styled.docx"

        MarkdownToWordConverter(font_name="Georgia").convert_content("Текст", output)

        document = Document(str(output))
        self.assertEqual(document.styles["Normal"].font.name, "Georgia")

    def test_preview_reports_warnings_without_writing_anything(self) -> None:
        source = self.root / "source.md"
        source.write_text("![diagram](missing.png)", encoding="utf-8")
        before = sorted(path.name for path in self.root.iterdir())

        warnings = MarkdownToWordConverter().preview_file(source)

        self.assertEqual(warnings, ["Image not found: missing.png"])
        self.assertEqual(sorted(path.name for path in self.root.iterdir()), before)

    def test_preview_of_clean_markdown_returns_no_warnings(self) -> None:
        warnings = MarkdownToWordConverter().preview_content("# Заголовок\n\nТекст.")

        self.assertEqual(warnings, [])

    def test_preview_of_a_missing_file_raises_conversion_error(self) -> None:
        with self.assertRaises(ConversionError):
            MarkdownToWordConverter().preview_file(self.root / "нет-такого.md")

    def test_allow_remote_images_defaults_to_true(self) -> None:
        # Guards the GUI-preservation guarantee: app.py builds this
        # converter without passing allow_remote_images, so it must default
        # to True (GfmDocxRenderer's own default) for the GUI to keep
        # fetching remote images unchanged. The MCP server is the only
        # caller that opts out, by passing allow_remote_images=False
        # explicitly (see mcp_server.py).
        self.assertIs(MarkdownToWordConverter().allow_remote_images, True)


class WordToMarkdownConverterTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmpdir.cleanup)
        self.root = Path(self._tmpdir.name)

    def test_headings_and_bold_runs_round_trip_into_markdown(self) -> None:
        docx_path = self.root / "source.docx"
        document = Document()
        document.add_heading("Раздел", level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run("жирный").bold = True
        document.save(str(docx_path))
        output = self.root / "source.md"

        warnings = WordToMarkdownConverter().convert_file(docx_path, output)

        text = output.read_text(encoding="utf-8")
        self.assertEqual(warnings, [])
        self.assertIn("## Раздел", text)
        self.assertIn("**жирный**", text)

    def test_inline_code_run_round_trips_with_the_original_wrapping(self) -> None:
        docx_path = self.root / "code.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("`code`")
        document.save(str(docx_path))
        output = self.root / "code.md"

        warnings = WordToMarkdownConverter().convert_file(docx_path, output)

        text = output.read_text(encoding="utf-8")
        self.assertEqual(warnings, [])
        # Ветка инлайн-кода оборачивает run.text ещё одной парой обратных
        # кавычек — так же, как в оригинале mdtoword/app.py.
        self.assertIn("``code``", text)

    def test_missing_source_raises_conversion_error(self) -> None:
        with self.assertRaises(ConversionError):
            WordToMarkdownConverter().convert_file(
                self.root / "нет-такого.docx", self.root / "out.md"
            )

    def test_heading_level_above_six_emits_plain_text_without_inline_markup(self) -> None:
        docx_path = self.root / "heading7.docx"
        document = Document()
        # В шаблоне python-docx уже есть скрытый (latent) стиль "Heading 7":
        # Word поддерживает заголовки уровней 1-9, а не только 1-6, так что
        # заводить стиль вручную не требуется.
        paragraph = document.add_paragraph(style="Heading 7")
        paragraph.add_run("жирный").bold = True
        document.save(str(docx_path))
        output = self.root / "heading7.md"

        warnings = WordToMarkdownConverter().convert_file(docx_path, output)

        text = output.read_text(encoding="utf-8")
        self.assertEqual(warnings, [])
        self.assertNotIn("#", text)
        self.assertNotIn("**", text)
        self.assertIn("жирный", text)


class CoreIsolationTests(unittest.TestCase):
    def test_importing_the_core_does_not_pull_in_pyqt(self) -> None:
        # Обязательно в подпроцессе: в общем прогоне tests/test_drop_queue.py
        # импортирует PyQt6, и проверка sys.modules в текущем процессе
        # проходила бы или падала в зависимости от порядка сбора тестов.
        repo_root = Path(__file__).resolve().parent.parent
        result = subprocess.run(
            [
                sys.executable,
                "-c",
                "import mdtoword.converters, sys; "
                "sys.exit(1 if 'PyQt6' in sys.modules else 0)",
            ],
            capture_output=True,
            text=True,
            cwd=repo_root,
        )

        self.assertEqual(result.returncode, 0, "ядро подтянуло PyQt6:\n" + result.stderr)


if __name__ == "__main__":
    unittest.main()
