import io
import unittest

from docx import Document
from docx.oxml.ns import nsmap as _nsmap
from docx.oxml.ns import qn

from mdtoword.latex_omml import (
    _ACCENTS,
    _ENVIRONMENT_ONLY,
    _ESCAPED,
    _LIMIT_OPERATORS,
    _NARY,
    _SPACING,
    _SYMBOLS,
    _UPRIGHT_FUNCTIONS,
    UnsupportedLatexError,
    latex_to_omml,
)

# python-docx registers a custom lxml element class per known tag, but it
# knows no `m:` (math) tags, so freshly built <m:...> elements come back as
# bare `lxml.etree._Element` without the `.xml` accessor this test module
# uses for inspection. Registering BaseOxmlElement as the namespace default
# upgrades every `m:` element so `.xml` works below. This is test-only
# convenience -- latex_omml.py itself never calls `.xml`, so the library has
# no business mutating python-docx's global registry to provide it.
try:
    from docx.oxml.parser import element_class_lookup as _element_class_lookup
    from docx.oxml.xmlchemy import BaseOxmlElement as _BaseOxmlElement

    _element_class_lookup.get_namespace(_nsmap["m"])[None] = _BaseOxmlElement
except (ImportError, KeyError, AttributeError):
    pass

MATH_NS = {"m": _nsmap["m"]}

# The exact set `_ENVIRONMENT_ONLY` is expected to have right now, written
# out by hand -- not derived from `_ENVIRONMENT_ONLY` at runtime, which
# would recreate the very self-referential hole the sweep test below would
# otherwise have on its own (deleting an entry here has to be a deliberate
# second edit).
EXPECTED_ENVIRONMENT_ONLY = (
    "matrix", "pmatrix", "bmatrix", "Bmatrix", "vmatrix", "Vmatrix",
    "array", "cases",
)


def xml_of(latex: str) -> str:
    return latex_to_omml(latex).xml


class LatexToOmmlTests(unittest.TestCase):
    def test_plain_run_marks_identifiers_italic_and_numbers_upright(self):
        xml = xml_of("x + 2")
        self.assertIn("<m:t>x</m:t>", xml)
        self.assertIn("<m:t>2</m:t>", xml)
        self.assertIn('<m:sty m:val="i"/>', xml)

    def test_fraction_builds_num_and_den(self):
        xml = xml_of(r"\frac{a+b}{2}")
        self.assertIn("<m:f>", xml)
        self.assertIn("<m:num>", xml)
        self.assertIn("<m:den>", xml)
        self.assertIn("<m:t>a</m:t>", xml)
        self.assertIn("<m:t>2</m:t>", xml)

    def test_superscript_and_subscript(self):
        self.assertIn("<m:sSup>", xml_of("x^2"))
        self.assertIn("<m:sSub>", xml_of("a_1"))
        self.assertIn("<m:sSubSup>", xml_of("x_i^2"))

    def test_square_root_with_and_without_degree(self):
        plain = xml_of(r"\sqrt{2}")
        self.assertIn("<m:rad>", plain)
        self.assertIn('<m:degHide m:val="1"/>', plain)
        cubic = xml_of(r"\sqrt[3]{8}")
        self.assertIn("<m:deg>", cubic)
        self.assertNotIn('<m:degHide m:val="1"/>', cubic)

    def test_greek_and_operators_map_to_unicode(self):
        self.assertIn("<m:t>α</m:t>", xml_of(r"\alpha"))
        self.assertIn("<m:t>∞</m:t>", xml_of(r"\infty"))
        self.assertIn("<m:t>≤</m:t>", xml_of(r"\leq"))
        self.assertIn("<m:t>×</m:t>", xml_of(r"\times"))

    def test_text_command_is_upright(self):
        xml = xml_of(r"\text{если}")
        self.assertIn("<m:t>если</m:t>", xml)
        self.assertIn('<m:nor m:val="1"/>', xml)

    def test_unsupported_command_raises_with_the_offending_token(self):
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\qedsymbol{x}")
        self.assertIn("qedsymbol", str(caught.exception))

    def test_unbalanced_braces_raise(self):
        with self.assertRaises(UnsupportedLatexError):
            latex_to_omml(r"\frac{a}{b")

    def test_unbraced_frac_argument_splits_multidigit_number(self):
        """`\\frac12x` means `\\frac{1}{2}x`, not `12/x` -- Finding 1."""
        element = latex_to_omml(r"\frac12x")
        children = list(element)
        self.assertEqual(len(children), 2)  # the fraction, then a separate "x" run
        fraction, trailing_run = children
        num = fraction.find("m:num", MATH_NS)
        den = fraction.find("m:den", MATH_NS)
        self.assertEqual([t.text for t in num.iter(qn("m:t"))], ["1"])
        self.assertEqual([t.text for t in den.iter(qn("m:t"))], ["2"])
        self.assertEqual([t.text for t in trailing_run.iter(qn("m:t"))], ["x"])

    def test_unbraced_superscript_splits_multidigit_number(self):
        """`x^12` superscripts only `1`, leaving a literal `2` outside -- Finding 1."""
        element = latex_to_omml("x^12")
        children = list(element)
        self.assertEqual(len(children), 2)  # the sSup, then a separate "2" run
        script, trailing_run = children
        sup = script.find("m:sup", MATH_NS)
        self.assertEqual([t.text for t in sup.iter(qn("m:t"))], ["1"])
        self.assertEqual([t.text for t in trailing_run.iter(qn("m:t"))], ["2"])

    def test_braced_number_arguments_still_take_the_whole_number(self):
        """Braced arguments must be unaffected by the unbraced-splitting fix."""
        frac_element = latex_to_omml(r"\frac{12}{x}")
        num = frac_element.find("m:f/m:num", MATH_NS)
        self.assertEqual([t.text for t in num.iter(qn("m:t"))], ["12"])
        sup_element = latex_to_omml("x^{10}")
        sup = sup_element.find("m:sSup/m:sup", MATH_NS)
        self.assertEqual([t.text for t in sup.iter(qn("m:t"))], ["10"])

    def test_every_environment_only_name_still_raises_as_a_bare_command(self):
        """Guards Finding 2: deleting `_ENVIRONMENT_ONLY` entries must not
        go unnoticed.

        `\\sum` degrading to a lone "Σ" with its limits silently dropped is
        exactly the failure mode this module exists to prevent, so every
        name that only works as an environment must raise -- naming the
        environment form to use instead -- rather than render as something
        plausible.  This sweeps the table itself, so an entry that is
        removed without being implemented has nowhere to hide.
        """
        self.assertTrue(
            _ENVIRONMENT_ONLY, "the environment-only table must stay guarded")
        # Deleting an entry from `_ENVIRONMENT_ONLY` above is not enough on
        # its own -- this table is iterated below, so the sweep can't notice
        # a name that is simply gone. Pinning the expected names against a
        # frozen literal forces a deliberate edit here too.
        self.assertEqual(
            sorted(_ENVIRONMENT_ONLY), sorted(EXPECTED_ENVIRONMENT_ONLY))
        for name in _ENVIRONMENT_ONLY:
            with self.subTest(command=name):
                with self.assertRaises(UnsupportedLatexError) as caught:
                    latex_to_omml("\\" + name)
                message = str(caught.exception)
                self.assertIn(name, message)
                # Not just the command name: the generic fallback message
                # also contains that (`Unsupported LaTeX command: \foo`), so
                # asserting only the name would not notice the
                # `_ENVIRONMENT_ONLY` enforcement block itself being
                # deleted. The `\begin{...}` suggestion only appears in the
                # dedicated `_ENVIRONMENT_ONLY` message.
                self.assertIn("\\begin{%s}" % name, message)

    def test_environment_only_names_are_disjoint_from_every_fallback_table(self):
        """`_ENVIRONMENT_ONLY` is checked before several branches for
        constructs this module already implements reach their own fallback
        tables -- not merely "before the symbol table" as the old comment in
        `_parse_command` claimed. If a name were ever added to
        `_ENVIRONMENT_ONLY` that also lived in one of those tables,
        whichever branch runs first would shadow the other, the same way a
        deleted entry used to fall through unnoticed (Finding 1). This pins
        the seven tables `_ENVIRONMENT_ONLY` must stay disjoint from."""
        fallback_tables = {
            "_NARY": _NARY,
            "_ACCENTS": _ACCENTS,
            "_LIMIT_OPERATORS": _LIMIT_OPERATORS,
            "_SYMBOLS": _SYMBOLS,
            "_SPACING": _SPACING,
            "_ESCAPED": _ESCAPED,
            "_UPRIGHT_FUNCTIONS": _UPRIGHT_FUNCTIONS,
        }
        for table_name, table in fallback_tables.items():
            with self.subTest(table=table_name):
                self.assertEqual(set(_ENVIRONMENT_ONLY) & set(table), set())

    def test_reserved_and_dangling_constructs_raise(self):
        """Constructs outside `_ENVIRONMENT_ONLY` that must still fail:
        unknown commands, and the halves of a pair used on their own."""
        must_raise = [
            (r"\qedsymbol", "qedsymbol"),
            (r"\begin{aligned} a \end{aligned}", "aligned"),
            (r"\left( x", "left"),
            (r"\right)", "right"),
            (r"\end{pmatrix}", "end"),
            (r"\begin{pmatrix} a & b", "end"),
            (r"\begin{pmatrix} a \end{bmatrix}", "bmatrix"),
            (r"\left\heartsuit x \right.", "heartsuit"),
            ("&", "&"),
        ]
        for latex, marker in must_raise:
            with self.subTest(latex=latex):
                with self.assertRaises(UnsupportedLatexError) as caught:
                    latex_to_omml(latex)
                self.assertIn(marker, str(caught.exception))

    def test_scripted_group_keeps_every_run_in_the_base(self):
        """`{a+b}^2` must nest all three runs -- a, +, b -- inside <m:e>,
        not just the last one (Finding 2)."""
        element = latex_to_omml("{a+b}^2")
        base = element.find("m:sSup/m:e", MATH_NS)
        self.assertIsNotNone(base)
        self.assertEqual(len(base), 3)

    def test_spacing_run_survives_docx_save_and_reopen(self):
        """A save/reopen round trip through python-docx's `remove_blank_text`
        parser must not eat the space run from `\\,` (Finding 2)."""
        document = Document()
        document.element.body.insert(0, latex_to_omml(r"a\,b"))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reopened = Document(buffer)
        texts = [t.text for t in reopened.element.body.iter(qn("m:t"))]
        self.assertEqual(texts, ["a", " ", "b"])

    def test_mathbf_is_bold_upright_boldsymbol_and_bm_are_bold_italic(self):
        """Finding 3: `\\mathbf` is bold upright; `\\boldsymbol`/`\\bm` are
        bold italic -- the previous implementation had these swapped."""
        mathbf_xml = xml_of(r"\mathbf{x}")
        self.assertIn('<m:sty m:val="b"/>', mathbf_xml)
        self.assertNotIn('<m:sty m:val="bi"/>', mathbf_xml)
        self.assertIn('<m:sty m:val="bi"/>', xml_of(r"\boldsymbol{x}"))
        self.assertIn('<m:sty m:val="bi"/>', xml_of(r"\bm{x}"))

    def test_square_root_with_empty_bracket_hides_degree(self):
        """`\\sqrt[]{x}` must behave like `\\sqrt{x}` -- Finding 5."""
        xml = xml_of(r"\sqrt[]{x}")
        self.assertIn('<m:degHide m:val="1"/>', xml)
        self.assertIn("<m:deg/>", xml)

    def test_no_run_combines_nor_with_sty(self):
        """OOXML's `CT_RPR` makes `<m:nor>` and `<m:sty>` a *choice*, not a
        sequence, so a run may carry one or the other but never both.

        `\\mathbf{x}` -- upright *and* bold -- is the case that hits it:
        emitting `<m:nor/><m:sty m:val="b"/>` is rejected by the ISO/IEC
        29500-4 schema. `<m:sty m:val="b">` already means bold *upright*, so
        the style value alone carries both facts and `<m:nor>` is redundant
        there. Swept over every construct that can produce a styled run.
        """
        formulas = [
            r"\mathbf{x}", r"\mathbf{ab}", r"\mathbf{2}",
            r"\mathbf{\text{ab}}", r"\mathbf{\alpha}",
            r"\text{hello}", r"\mathrm{d}", r"\operatorname{sgn}",
            r"\boldsymbol{v}", r"\bm{w}", r"\mathit{y}",
            r"\mathbf{\frac{a}{b}}", r"\mathbf{\sin x}",
        ]
        for latex in formulas:
            with self.subTest(latex=latex):
                for properties in latex_to_omml(latex).iter(qn("m:rPr")):
                    present = _tags(properties)
                    self.assertFalse(
                        "nor" in present and "sty" in present,
                        f"<m:rPr> carries both nor and sty: {present}",
                    )


def _tags(element) -> list:
    """Local names of an element's direct children, in document order."""
    return [child.tag.split("}")[-1] for child in element]


def _texts(element) -> list:
    """Every <m:t> under `element`, in document order."""
    return [t.text for t in element.iter(qn("m:t"))]


class BigConstructTests(unittest.TestCase):
    """Task 3: n-ary operators, limits, delimiters, accents, matrices.

    These assert on element structure rather than substrings: an <m:nary>
    whose limits landed in the body would satisfy every `assertIn` a
    grep-style test can make, yet render wrongly in Word.
    """

    def test_nary_sum_carries_limits(self):
        element = latex_to_omml(r"\sum_{i=1}^{n} i")
        nary = element.find("m:nary", MATH_NS)
        self.assertIsNotNone(nary)
        self.assertEqual(_tags(nary), ["naryPr", "sub", "sup", "e"])
        self.assertEqual(
            nary.find("m:naryPr/m:chr", MATH_NS).get(qn("m:val")), "∑")
        self.assertEqual(
            nary.find("m:naryPr/m:limLoc", MATH_NS).get(qn("m:val")), "undOvr")
        self.assertEqual(_texts(nary.find("m:sub", MATH_NS)), ["i", "=", "1"])
        self.assertEqual(_texts(nary.find("m:sup", MATH_NS)), ["n"])
        self.assertEqual(_texts(nary.find("m:e", MATH_NS)), ["i"])

    def test_nary_without_limits_hides_the_empty_boxes(self):
        nary = latex_to_omml(r"\prod a").find("m:nary", MATH_NS)
        self.assertEqual(
            nary.find("m:naryPr/m:chr", MATH_NS).get(qn("m:val")), "∏")
        self.assertEqual(
            nary.find("m:naryPr/m:subHide", MATH_NS).get(qn("m:val")), "1")
        self.assertEqual(
            nary.find("m:naryPr/m:supHide", MATH_NS).get(qn("m:val")), "1")
        self.assertEqual(_texts(nary.find("m:e", MATH_NS)), ["a"])

    def test_integral_uses_its_own_character_and_side_limits(self):
        nary = latex_to_omml(r"\int_0^1 x\,dx").find("m:nary", MATH_NS)
        self.assertEqual(
            nary.find("m:naryPr/m:chr", MATH_NS).get(qn("m:val")), "∫")
        self.assertEqual(
            nary.find("m:naryPr/m:limLoc", MATH_NS).get(qn("m:val")), "subSup")
        self.assertEqual(_texts(nary.find("m:sub", MATH_NS)), ["0"])
        self.assertEqual(_texts(nary.find("m:sup", MATH_NS)), ["1"])
        self.assertEqual(_texts(nary.find("m:e", MATH_NS)), ["x", " ", "d", "x"])

    def test_every_integral_sign_keeps_limits_beside_it(self):
        """∬, ∭ and ∮ are integrals too -- not just the plain ∫."""
        for latex, character in ((r"\iint x", "∬"), (r"\iiint x", "∭"),
                                 (r"\oint x", "∮")):
            with self.subTest(latex=latex):
                properties = latex_to_omml(latex).find("m:nary/m:naryPr", MATH_NS)
                self.assertEqual(
                    properties.find("m:chr", MATH_NS).get(qn("m:val")), character)
                self.assertEqual(
                    properties.find("m:limLoc", MATH_NS).get(qn("m:val")), "subSup")

    def test_limit_uses_lim_low(self):
        element = latex_to_omml(r"\lim_{x \to 0} f(x)")
        limit = element.find("m:limLow", MATH_NS)
        self.assertIsNotNone(limit)
        self.assertEqual(_tags(limit), ["e", "lim"])
        self.assertEqual(_texts(limit.find("m:e", MATH_NS)), ["lim"])
        self.assertEqual(_texts(limit.find("m:lim", MATH_NS)), ["x", "→", "0"])
        # `f(x)` is the limit's operand in prose, but it stays a sibling in
        # OMML -- it must not be swallowed into <m:lim>.
        self.assertEqual(
            _texts(element), ["lim", "x", "→", "0", "f", "(", "x", ")"])

    def test_limsup_and_liminf_produce_lim_low_with_the_right_text(self):
        """`\\limsup`/`\\liminf` left the unimplemented table and were
        mapping to the upright run text "lim sup"/"lim inf" -- new behaviour
        that gained zero coverage, the exact pattern these guards exist to
        catch. `\\limsup_{n} a` must produce a `limLow` whose `<m:e>` holds
        the text "lim sup", not a bare symbol with its limit dropped."""
        element = latex_to_omml(r"\limsup_{n} a")
        limit = element.find("m:limLow", MATH_NS)
        self.assertIsNotNone(limit)
        self.assertEqual(_tags(limit), ["e", "lim"])
        self.assertEqual(_texts(limit.find("m:e", MATH_NS)), ["lim sup"])
        self.assertEqual(_texts(limit.find("m:lim", MATH_NS)), ["n"])
        self.assertEqual(_texts(element), ["lim sup", "n", "a"])

        liminf_limit = latex_to_omml(r"\liminf_{n} a").find(
            "m:limLow", MATH_NS)
        self.assertIsNotNone(liminf_limit)
        self.assertEqual(_texts(liminf_limit.find("m:e", MATH_NS)), ["lim inf"])

    def test_bare_limit_without_a_subscript_is_just_an_upright_run(self):
        element = latex_to_omml(r"\lim f")
        self.assertIsNone(element.find("m:limLow", MATH_NS))
        self.assertEqual(_texts(element), ["lim", "f"])
        self.assertIn('<m:nor m:val="1"/>', element.xml)

    def test_left_right_delimiters_wrap_the_whole_fraction(self):
        element = latex_to_omml(r"\left( \frac{a}{b} \right)")
        delimiter = element.find("m:d", MATH_NS)
        self.assertIsNotNone(delimiter)
        self.assertEqual(
            delimiter.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), ")")
        body = delimiter.find("m:e", MATH_NS)
        self.assertEqual(len(body), 1)
        self.assertEqual(body[0].tag, qn("m:f"))

    def test_left_right_supports_dot_braces_and_bars(self):
        pairs = {
            r"\left[ x \right]": ("[", "]"),
            r"\left\{ x \right\}": ("{", "}"),
            r"\left| x \right|": ("|", "|"),
            r"\left\| x \right\|": ("‖", "‖"),
            r"\left\langle x \right\rangle": ("⟨", "⟩"),
            r"\left. x \right)": ("", ")"),
        }
        for latex, (begin, end) in pairs.items():
            with self.subTest(latex=latex):
                properties = latex_to_omml(latex).find("m:d/m:dPr", MATH_NS)
                self.assertEqual(
                    properties.find("m:begChr", MATH_NS).get(qn("m:val")), begin)
                self.assertEqual(
                    properties.find("m:endChr", MATH_NS).get(qn("m:val")), end)

    def test_nested_left_right_pairs_match_innermost_first(self):
        outer = latex_to_omml(r"\left[ \left( a \right) + b \right]").find(
            "m:d", MATH_NS)
        self.assertEqual(
            outer.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "[")
        body = outer.find("m:e", MATH_NS)
        self.assertEqual([child.tag.split("}")[-1] for child in body],
                         ["d", "r", "r"])
        self.assertEqual(
            body.find("m:d/m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")

    def test_accents_and_overline(self):
        accent = latex_to_omml(r"\hat{x}").find("m:acc", MATH_NS)
        self.assertEqual(_tags(accent), ["accPr", "e"])
        self.assertEqual(
            accent.find("m:accPr/m:chr", MATH_NS).get(qn("m:val")), "̂")
        self.assertEqual(_texts(accent.find("m:e", MATH_NS)), ["x"])

        bar = latex_to_omml(r"\overline{AB}").find("m:bar", MATH_NS)
        self.assertEqual(_tags(bar), ["barPr", "e"])
        self.assertEqual(
            bar.find("m:barPr/m:pos", MATH_NS).get(qn("m:val")), "top")
        self.assertEqual(_texts(bar.find("m:e", MATH_NS)), ["A", "B"])

        self.assertEqual(
            latex_to_omml(r"\underline{x}")
            .find("m:bar/m:barPr/m:pos", MATH_NS).get(qn("m:val")), "bot")
        self.assertIn('<m:chr m:val="⃗"/>', xml_of(r"\vec{v}"))

    def test_accent_still_accepts_its_own_scripts(self):
        """`\\hat{x}^2` superscripts the accented base, not a bare `x`."""
        element = latex_to_omml(r"\hat{x}^2")
        base = element.find("m:sSup/m:e", MATH_NS)
        self.assertEqual(len(base), 1)
        self.assertEqual(base[0].tag, qn("m:acc"))
        self.assertEqual(_texts(element.find("m:sSup/m:sup", MATH_NS)), ["2"])

    def test_binomial_is_a_barless_fraction_in_parentheses(self):
        delimiter = latex_to_omml(r"\binom{n}{k}").find("m:d", MATH_NS)
        self.assertEqual(
            delimiter.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), ")")
        body = delimiter.find("m:e", MATH_NS)
        self.assertEqual(len(body), 1)
        fraction = body[0]
        self.assertEqual(fraction.tag, qn("m:f"))
        self.assertEqual(
            fraction.find("m:fPr/m:type", MATH_NS).get(qn("m:val")), "noBar")
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["n"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["k"])

    def test_pmatrix_builds_a_matrix_in_parentheses(self):
        element = latex_to_omml(r"\begin{pmatrix} a & b \\ c & d \end{pmatrix}")
        delimiter = element.find("m:d", MATH_NS)
        self.assertEqual(
            delimiter.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), ")")
        rows = delimiter.findall("m:e/m:m/m:mr", MATH_NS)
        self.assertEqual(len(rows), 2)
        self.assertEqual(
            [len(row.findall("m:e", MATH_NS)) for row in rows], [2, 2])
        self.assertEqual(
            [_texts(cell) for row in rows
             for cell in row.findall("m:e", MATH_NS)],
            [["a"], ["b"], ["c"], ["d"]])

    def test_matrix_flavours_pick_their_own_fences(self):
        fences = {
            "matrix": None,
            "pmatrix": ("(", ")"),
            "bmatrix": ("[", "]"),
            "Bmatrix": ("{", "}"),
            "vmatrix": ("|", "|"),
            "Vmatrix": ("‖", "‖"),
        }
        for name, pair in fences.items():
            with self.subTest(environment=name):
                element = latex_to_omml(
                    "\\begin{%s} a & b \\\\ c & d \\end{%s}" % (name, name))
                if pair is None:
                    self.assertIsNone(element.find("m:d", MATH_NS))
                    self.assertIsNotNone(element.find("m:m", MATH_NS))
                    continue
                properties = element.find("m:d/m:dPr", MATH_NS)
                self.assertEqual(
                    properties.find("m:begChr", MATH_NS).get(qn("m:val")), pair[0])
                self.assertEqual(
                    properties.find("m:endChr", MATH_NS).get(qn("m:val")), pair[1])

    def test_cases_environment(self):
        element = latex_to_omml(
            r"\begin{cases} x & x > 0 \\ -x & x \leq 0 \end{cases}")
        delimiter = element.find("m:d", MATH_NS)
        self.assertEqual(
            delimiter.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "{")
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), "")
        rows = delimiter.findall("m:e/m:m/m:mr", MATH_NS)
        self.assertEqual(len(rows), 2)
        self.assertEqual(
            [len(row.findall("m:e", MATH_NS)) for row in rows], [2, 2])
        self.assertEqual(
            [_texts(cell) for cell in rows[1].findall("m:e", MATH_NS)],
            [["-", "x"], ["x", "≤", "0"]])

    def test_ragged_rows_are_padded_so_word_sees_a_rectangle(self):
        rows = latex_to_omml(
            r"\begin{cases} a & b \\ c \end{cases}"
        ).findall("m:d/m:e/m:m/m:mr", MATH_NS)
        self.assertEqual(
            [len(row.findall("m:e", MATH_NS)) for row in rows], [2, 2])
        self.assertEqual(_texts(rows[1].findall("m:e", MATH_NS)[1]), [])

    def test_trailing_row_separator_does_not_add_an_empty_row(self):
        rows = latex_to_omml(
            r"\begin{pmatrix} a & b \\ c & d \\ \end{pmatrix}"
        ).findall("m:d/m:e/m:m/m:mr", MATH_NS)
        self.assertEqual(len(rows), 2)

    def test_nary_body_stops_at_the_enclosing_construct(self):
        """A sum's body runs to the end of its group -- but not past the
        `\\right`, the `&` or the `}` that closes it."""
        delimiter = latex_to_omml(
            r"\left( \sum_{i} a_i \right) + b").find("m:d", MATH_NS)
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), ")")
        self.assertEqual(_texts(delimiter.find("m:e", MATH_NS)), ["i", "a", "i"])

        rows = latex_to_omml(
            r"\begin{pmatrix} \sum_i a_i & b \end{pmatrix}"
        ).findall("m:d/m:e/m:m/m:mr", MATH_NS)
        self.assertEqual(
            [len(row.findall("m:e", MATH_NS)) for row in rows], [2])
        self.assertEqual(_texts(rows[0].findall("m:e", MATH_NS)[1]), ["b"])

        group = latex_to_omml(r"{\sum_i a_i}b")
        self.assertEqual(_tags(group), ["nary", "r"])
        self.assertEqual(_texts(group[1]), ["b"])

    def test_big_constructs_survive_docx_save_and_reopen(self):
        """The shapes must come back intact through a real save/open cycle,
        including python-docx's blank-text-stripping parser."""
        document = Document()
        formulas = [
            r"\sum_{i=1}^{n} i",
            r"\int_0^1 x\,dx",
            r"\begin{pmatrix} a & b \\ c & d \end{pmatrix}",
            r"\begin{cases} x & x > 0 \\ -x & x \leq 0 \end{cases}",
        ]
        for offset, formula in enumerate(formulas):
            document.element.body.insert(offset, latex_to_omml(formula))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        body = Document(buffer).element.body

        naries = body.findall("m:oMath/m:nary", MATH_NS)
        self.assertEqual(
            [n.find("m:naryPr/m:chr", MATH_NS).get(qn("m:val")) for n in naries],
            ["∑", "∫"])
        self.assertEqual(_texts(naries[0].find("m:sub", MATH_NS)), ["i", "=", "1"])
        self.assertEqual(_texts(naries[0].find("m:e", MATH_NS)), ["i"])
        self.assertEqual(_texts(naries[1].find("m:e", MATH_NS)),
                         ["x", " ", "d", "x"])

        matrices = body.findall("m:oMath/m:d/m:e/m:m", MATH_NS)
        self.assertEqual(len(matrices), 2)
        self.assertEqual(
            [len(m.findall("m:mr", MATH_NS)) for m in matrices], [2, 2])
        self.assertEqual(_texts(matrices[0]), ["a", "b", "c", "d"])
        self.assertEqual(
            [d.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val"))
             for d in body.findall("m:oMath/m:d", MATH_NS)], ["(", "{"])

    def test_every_nary_and_accent_entry_converts_without_raising(self):
        """13 of these -- \\coprod, \\bigoplus, \\bigotimes, \\bigvee,
        \\bigwedge, \\widehat, \\widetilde, \\dot, \\ddot, \\acute, \\grave,
        \\check, \\breve -- had zero coverage: new-behaviour-with-no-test is
        exactly the pattern the `_ENVIRONMENT_ONLY` guards exist to catch,
        and these tables are no different just because they are already
        implemented."""
        for name, character in _NARY.items():
            with self.subTest(command=name):
                nary = latex_to_omml("\\%s x" % name).find("m:nary", MATH_NS)
                self.assertIsNotNone(nary)
                self.assertEqual(
                    nary.find("m:naryPr/m:chr", MATH_NS).get(qn("m:val")),
                    character)

        for name, character in _ACCENTS.items():
            with self.subTest(command=name):
                accent = latex_to_omml("\\%s{x}" % name).find("m:acc", MATH_NS)
                self.assertIsNotNone(accent)
                self.assertEqual(
                    accent.find("m:accPr/m:chr", MATH_NS).get(qn("m:val")),
                    character)


class InfixCommandTests(unittest.TestCase):
    r"""``\over``, ``\atop`` and ``\choose``: TeX's infix fraction builders.

    Each one splits the group it appears in -- everything to its left is the
    numerator, everything to its right the denominator -- which is why they
    cannot be handled in ``_parse_command`` like a prefix command.
    """

    def test_over_builds_a_plain_fraction(self):
        fraction = latex_to_omml(r"{a + b \over c}").find("m:f", MATH_NS)
        self.assertIsNotNone(fraction)
        self.assertIsNone(fraction.find("m:fPr/m:type", MATH_NS))
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["a", "+", "b"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["c"])

    def test_atop_builds_a_barless_fraction(self):
        fraction = latex_to_omml(r"{n \atop k}").find("m:f", MATH_NS)
        self.assertEqual(
            fraction.find("m:fPr/m:type", MATH_NS).get(qn("m:val")), "noBar")
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["n"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["k"])
        self.assertIsNone(
            latex_to_omml(r"{n \atop k}").find("m:d", MATH_NS),
            r"\atop stacks without fences -- those belong to \choose",
        )

    def test_choose_is_a_barless_fraction_in_parentheses(self):
        r"""``{n \choose k}`` must match ``\binom{n}{k}`` exactly."""
        delimiter = latex_to_omml(r"{n \choose k}").find("m:d", MATH_NS)
        self.assertEqual(
            delimiter.find("m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")
        self.assertEqual(
            delimiter.find("m:dPr/m:endChr", MATH_NS).get(qn("m:val")), ")")
        fraction = delimiter.find("m:e/m:f", MATH_NS)
        self.assertEqual(
            fraction.find("m:fPr/m:type", MATH_NS).get(qn("m:val")), "noBar")
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["n"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["k"])

    def test_infix_without_braces_splits_the_whole_formula(self):
        fraction = latex_to_omml(r"1 \over 2").find("m:f", MATH_NS)
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["1"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["2"])

    def test_infix_stops_at_the_group_that_encloses_it(self):
        r"""``x + {a \over b} + y`` divides only inside the braces."""
        element = latex_to_omml(r"x + {a \over b} + y")
        self.assertEqual(
            _tags(element), ["r", "r", "f", "r", "r"])
        fraction = element.find("m:f", MATH_NS)
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["a"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["b"])

    def test_infix_inside_a_matrix_cell_stops_at_the_cell(self):
        rows = latex_to_omml(
            r"\begin{pmatrix} a \over b & c \end{pmatrix}"
        ).findall("m:d/m:e/m:m/m:mr", MATH_NS)
        cells = rows[0].findall("m:e", MATH_NS)
        self.assertEqual(len(cells), 2)
        self.assertEqual(_tags(cells[0]), ["f"])
        self.assertEqual(_texts(cells[1]), ["c"])

    def test_two_infix_commands_in_one_group_raise(self):
        r"""TeX itself rejects ``a \over b \over c`` as ambiguous; guessing
        an association here would silently produce one of two readings."""
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"a \over b \over c")
        self.assertIn("over", str(caught.exception))

    def test_nesting_infix_through_braces_is_allowed(self):
        r"""``{{a \over b} \over c}`` is unambiguous and must convert."""
        outer = latex_to_omml(r"{{a \over b} \over c}").find("m:f", MATH_NS)
        self.assertEqual(_tags(outer.find("m:num", MATH_NS)), ["f"])
        self.assertEqual(_texts(outer.find("m:den", MATH_NS)), ["c"])

    def test_infix_with_nothing_after_it_still_builds_an_empty_denominator(self):
        fraction = latex_to_omml(r"a \over").find("m:f", MATH_NS)
        self.assertEqual(_texts(fraction.find("m:num", MATH_NS)), ["a"])
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), [])

    def test_infix_used_as_an_argument_reports_its_two_sides(self):
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\frac\over x")
        self.assertIn("over", str(caught.exception))


class LineBreakTests(unittest.TestCase):
    r"""``\\`` outside a matrix: an OMML equation array (``<m:eqArr>``)."""

    def test_line_break_builds_an_equation_array(self):
        element = latex_to_omml(r"a = b \\ c = d")
        array = element.find("m:eqArr", MATH_NS)
        self.assertIsNotNone(array)
        lines = array.findall("m:e", MATH_NS)
        self.assertEqual(len(lines), 2)
        self.assertEqual(_texts(lines[0]), ["a", "=", "b"])
        self.assertEqual(_texts(lines[1]), ["c", "=", "d"])

    def test_a_formula_without_a_line_break_stays_flat(self):
        self.assertIsNone(latex_to_omml("a = b").find("m:eqArr", MATH_NS))

    def test_trailing_line_break_does_not_add_an_empty_line(self):
        element = latex_to_omml(r"a \\ b \\")
        lines = element.find("m:eqArr", MATH_NS).findall("m:e", MATH_NS)
        self.assertEqual(len(lines), 2)

    def test_a_lone_line_break_produces_nothing_rather_than_raising(self):
        r"""``\\`` on its own is an empty line, not an error: the whole
        point of implementing it is that it no longer fails loudly."""
        element = latex_to_omml("\\\\")
        self.assertEqual(list(element), [])

    def test_line_breaks_nest_inside_a_group(self):
        element = latex_to_omml(r"\left\{ a \\ b \right.")
        lines = element.findall("m:d/m:e/m:eqArr/m:e", MATH_NS)
        self.assertEqual([_texts(line) for line in lines], [["a"], ["b"]])

    def test_a_line_break_ends_an_n_ary_operator_body(self):
        r"""``\sum_i a_i \\ b`` must leave ``b`` on the next line, not
        swallow it into the sum's operand."""
        element = latex_to_omml(r"\sum_i a_i \\ b")
        lines = element.findall("m:eqArr/m:e", MATH_NS)
        self.assertEqual(len(lines), 2)
        self.assertEqual(_tags(lines[0]), ["nary"])
        self.assertEqual(_texts(lines[0].find("m:nary/m:e", MATH_NS)), ["a", "i"])
        self.assertEqual(_texts(lines[1]), ["b"])

    def test_matrix_rows_are_still_rows_not_an_equation_array(self):
        element = latex_to_omml(r"\begin{pmatrix} a \\ b \end{pmatrix}")
        self.assertIsNone(element.find(".//m:eqArr", MATH_NS))
        self.assertEqual(len(element.findall("m:d/m:e/m:m/m:mr", MATH_NS)), 2)


class AlignmentPointTests(unittest.TestCase):
    r"""``&`` between the lines of a multi-line formula: an OMML alignment
    point, spelled ``<m:aln/>`` inside the ``<m:rPr>`` of the run that
    starts the aligned segment -- which is how Word itself writes it.

    This is what lets ``align`` stay ONE Word equation with its ``=`` signs
    lined up, instead of being cut into one centred paragraph per line.
    """

    @staticmethod
    def _aligned_runs(line):
        """Every run in `line` carrying an alignment point, with its text."""
        return [
            "".join(t.text or "" for t in run.iter(qn("m:t")))
            for run in line.iter(qn("m:r"))
            if run.find("m:rPr/m:aln", MATH_NS) is not None
        ]

    def test_ampersand_marks_the_run_that_follows_it(self):
        r"""In ``a &= b`` the ``&`` sits before ``=``, so the ``=`` run is
        the alignment point -- exactly what lines the equals signs up."""
        element = latex_to_omml(r"a &= b \\ c &= d")
        lines = element.findall("m:eqArr/m:e", MATH_NS)
        self.assertEqual(len(lines), 2)
        for line in lines:
            self.assertEqual(self._aligned_runs(line), ["="])
        self.assertEqual(
            [_texts(line) for line in lines],
            [["a", "=", "b"], ["c", "=", "d"]])

    def test_alignment_point_does_not_add_or_drop_any_text(self):
        r"""The marker is formatting on an existing run, not a new glyph."""
        self.assertEqual(
            _texts(latex_to_omml(r"a &= b \\ c &= d")),
            _texts(latex_to_omml(r"a = b \\ c = d")))

    def test_several_alignment_points_on_one_line(self):
        r"""``eqnarray``'s ``a &=& b`` has two, and both must survive."""
        line = latex_to_omml(r"a &=& b \\ c &=& d").findall(
            "m:eqArr/m:e", MATH_NS)[0]
        self.assertEqual(self._aligned_runs(line), ["=", "b"])

    def test_alignment_point_before_a_non_run_gets_its_own_marker(self):
        r"""``a &\frac{1}{2}`` puts the point before a fraction, which has
        no ``<m:rPr>`` to carry it, so an empty marker run stands in."""
        line = latex_to_omml(
            r"a &\frac{1}{2} \\ b &\frac{3}{4}"
        ).findall("m:eqArr/m:e", MATH_NS)[0]
        self.assertEqual(_tags(line), ["r", "r", "f"])
        self.assertEqual(self._aligned_runs(line), [""])
        # The marker's own <m:t> is empty, so it contributes no glyph.
        self.assertEqual(_texts(line), ["a", "", "1", "2"])

    def test_alignment_point_at_the_end_of_a_line_is_still_recorded(self):
        line = latex_to_omml(r"a & \\ b & c").findall(
            "m:eqArr/m:e", MATH_NS)[0]
        self.assertEqual(self._aligned_runs(line), [""])

    def test_an_alignment_point_ends_an_n_ary_operator_body(self):
        r"""``\sum_i a_i &= b`` must leave ``a_i`` as the sum's operand and
        let ``&`` start the next aligned segment.

        Without this the sum swallows the ``&`` into its own body, where --
        having no line break of its own to pair with -- it is rejected as a
        stray ampersand and the whole formula fails to convert. Found by
        validating against the OOXML schema, not by an earlier test.
        """
        element = latex_to_omml(r"\sum_i a_i &= b \\ c &= d")
        lines = element.findall("m:eqArr/m:e", MATH_NS)
        self.assertEqual(len(lines), 2)
        self.assertEqual(_tags(lines[0]), ["nary", "r", "r"])
        nary = lines[0].find("m:nary", MATH_NS)
        self.assertEqual(_texts(nary.find("m:e", MATH_NS)), ["a", "i"])
        self.assertEqual(self._aligned_runs(lines[0]), ["="])

    def test_an_alignment_point_ends_an_infix_denominator(self):
        r"""Same rule for ``\over``: ``{a \over b} &= c`` style input must
        not pull the ``&`` into the denominator."""
        element = latex_to_omml(r"a \over b &= c \\ d &= e")
        lines = element.findall("m:eqArr/m:e", MATH_NS)
        fraction = lines[0].find("m:f", MATH_NS)
        self.assertEqual(_texts(fraction.find("m:den", MATH_NS)), ["b"])
        self.assertEqual(self._aligned_runs(lines[0]), ["="])

    def test_stray_ampersand_in_a_single_line_formula_still_raises(self):
        r"""Without a ``\\`` there is nothing to align against, so ``&`` is
        far more likely a literal ampersand the user forgot to escape.
        Keeping the loud error is what stops ``Tom & Jerry`` inside ``$…$``
        from turning into an invisible alignment marker."""
        for latex in ("a & b", "&", r"\text{x} & y"):
            with self.subTest(latex=latex):
                with self.assertRaises(UnsupportedLatexError) as caught:
                    latex_to_omml(latex)
                self.assertIn("&", str(caught.exception))

    def test_matrix_ampersands_stay_column_separators(self):
        r"""A matrix consumes ``&`` as a cell break long before the
        alignment logic sees it -- no marker may leak into one."""
        element = latex_to_omml(r"\begin{pmatrix} a & b \\ c & d \end{pmatrix}")
        self.assertIsNone(element.find(".//m:aln", MATH_NS))
        self.assertEqual(
            len(element.findall("m:d/m:e/m:m/m:mr/m:e", MATH_NS)), 4)

    def test_array_ampersands_stay_column_separators(self):
        element = latex_to_omml(
            r"\begin{array}{cc} a & b \\ c & d \end{array}")
        self.assertIsNone(element.find(".//m:aln", MATH_NS))

    def test_alignment_point_survives_a_docx_save_and_reopen(self):
        document = Document()
        document.element.body.insert(0, latex_to_omml(r"a &= b \\ c &= d"))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        lines = Document(buffer).element.body.findall(
            "m:oMath/m:eqArr/m:e", MATH_NS)
        self.assertEqual(len(lines), 2)
        for line in lines:
            self.assertEqual(self._aligned_runs(line), ["="])


class ArrayAndSubstackTests(unittest.TestCase):
    r"""``\begin{array}{...}`` and ``\substack{...}``."""

    def test_array_builds_an_unfenced_matrix(self):
        element = latex_to_omml(
            r"\begin{array}{cc} a & b \\ c & d \end{array}")
        self.assertIsNone(element.find("m:d", MATH_NS))
        rows = element.findall("m:m/m:mr", MATH_NS)
        self.assertEqual(len(rows), 2)
        self.assertEqual(
            [_texts(cell) for row in rows
             for cell in row.findall("m:e", MATH_NS)],
            [["a"], ["b"], ["c"], ["d"]])

    def test_array_column_specification_sets_each_column_justification(self):
        matrix = latex_to_omml(
            r"\begin{array}{lcr} a & b & c \end{array}").find("m:m", MATH_NS)
        columns = matrix.findall("m:mPr/m:mcs/m:mc", MATH_NS)
        self.assertEqual(
            [c.find("m:mcPr/m:mcJc", MATH_NS).get(qn("m:val")) for c in columns],
            ["left", "center", "right"])
        self.assertEqual(
            [c.find("m:mcPr/m:count", MATH_NS).get(qn("m:val")) for c in columns],
            ["1", "1", "1"])

    def test_array_properties_come_before_the_rows(self):
        """OOXML requires <m:mPr> first; Word rejects the file otherwise."""
        matrix = latex_to_omml(
            r"\begin{array}{c} a \end{array}").find("m:m", MATH_NS)
        self.assertEqual(_tags(matrix)[0], "mPr")

    def test_array_fenced_by_left_right_keeps_both(self):
        element = latex_to_omml(
            r"\left( \begin{array}{cc} a & b \end{array} \right)")
        self.assertEqual(
            element.find("m:d/m:dPr/m:begChr", MATH_NS).get(qn("m:val")), "(")
        self.assertIsNotNone(element.find("m:d/m:e/m:m", MATH_NS))

    def test_array_without_a_column_specification_says_so(self):
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\begin{array} a \end{array}")
        self.assertIn("column specification", str(caught.exception))

    def test_array_vertical_rule_is_rejected_rather_than_dropped(self):
        r"""OMML has no vertical rule inside a matrix, so ``{c|c}`` cannot
        be honoured -- silently dropping the rule would turn an augmented
        matrix into an ordinary one."""
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\begin{array}{c|c} a & b \end{array}")
        message = str(caught.exception)
        self.assertIn("|", message)
        self.assertIn("array", message)

    def test_array_with_more_columns_than_declared_is_rejected(self):
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\begin{array}{c} a & b \end{array}")
        self.assertIn("column", str(caught.exception))

    def test_array_padded_column_from_the_specification_stays_empty(self):
        rows = latex_to_omml(
            r"\begin{array}{ccc} a & b \end{array}").findall("m:m/m:mr", MATH_NS)
        cells = rows[0].findall("m:e", MATH_NS)
        self.assertEqual(len(cells), 3)
        self.assertEqual(_texts(cells[2]), [])

    def test_substack_stacks_its_lines_in_a_single_column(self):
        matrix = latex_to_omml(
            r"\substack{i < j \\ i \in S}").find("m:m", MATH_NS)
        self.assertIsNotNone(matrix)
        rows = matrix.findall("m:mr", MATH_NS)
        self.assertEqual(len(rows), 2)
        self.assertEqual(
            [len(row.findall("m:e", MATH_NS)) for row in rows], [1, 1])
        self.assertEqual(_texts(rows[0]), ["i", "<", "j"])
        self.assertEqual(_texts(rows[1]), ["i", "∈", "S"])

    def test_substack_serves_as_an_n_ary_limit(self):
        r"""Its whole reason to exist: ``\sum_{\substack{...}}``."""
        nary = latex_to_omml(
            r"\sum_{\substack{i = 1 \\ i \neq j}} a_i").find("m:nary", MATH_NS)
        self.assertEqual(
            len(nary.findall("m:sub/m:m/m:mr", MATH_NS)), 2)
        self.assertEqual(_texts(nary.find("m:e", MATH_NS)), ["a", "i"])

    def test_substack_without_a_brace_group_says_so(self):
        with self.assertRaises(UnsupportedLatexError) as caught:
            latex_to_omml(r"\substack x")
        self.assertIn("substack", str(caught.exception))

    def test_new_constructs_survive_a_docx_save_and_reopen(self):
        """The shapes must come back intact through a real save/open cycle,
        including python-docx's blank-text-stripping parser."""
        document = Document()
        formulas = [
            r"\begin{array}{lc} a & b \\ c & d \end{array}",
            r"{n \choose k}",
            r"x = 1 \\ y = 2",
            r"\sum_{\substack{i \\ j}} a",
        ]
        for offset, formula in enumerate(formulas):
            document.element.body.insert(offset, latex_to_omml(formula))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        body = Document(buffer).element.body

        array = body.find("m:oMath/m:m", MATH_NS)
        self.assertEqual(
            [c.find("m:mcPr/m:mcJc", MATH_NS).get(qn("m:val"))
             for c in array.findall("m:mPr/m:mcs/m:mc", MATH_NS)],
            ["left", "center"])
        self.assertEqual(_texts(array), ["a", "b", "c", "d"])

        binomial = body.find("m:oMath/m:d/m:e/m:f", MATH_NS)
        self.assertEqual(
            binomial.find("m:fPr/m:type", MATH_NS).get(qn("m:val")), "noBar")

        lines = body.findall("m:oMath/m:eqArr/m:e", MATH_NS)
        self.assertEqual(
            [_texts(line) for line in lines], [["x", "=", "1"], ["y", "=", "2"]])

        self.assertEqual(
            len(body.findall("m:oMath/m:nary/m:sub/m:m/m:mr", MATH_NS)), 2)


if __name__ == "__main__":
    unittest.main()
