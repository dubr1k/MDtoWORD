"""Ядро конвертации Markdown ↔ Word.

Модуль намеренно свободен от PyQt6: им пользуются и GUI (``mdtoword.app``),
и MCP-сервер (``mdtoword.mcp_server``), а последний обязан работать без
графической подсистемы.

Контракт: успешная конвертация возвращает список необязательных предупреждений,
неуспешная — бросает :class:`ConversionError`. Сообщения об ошибках намеренно
не локализованы: язык подставляет потребитель, у GUI для этого есть словарь
переводов, а агенту локаль не нужна вовсе.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from .gfm_renderer import GfmDocxRenderer


class ConversionError(Exception):
    """Конвертация не удалась. Сообщение пригодно для показа без перевода."""


class MarkdownToWordConverter:
    """Конвертирует GFM-разметку в документ Word."""

    def __init__(
        self,
        font_name: str = "Times New Roman",
        font_size: Pt = Pt(12),
        footnotes_heading: str = "Footnotes",
    ) -> None:
        self.default_font_name = font_name
        self.default_font_size = font_size
        self.footnotes_heading = footnotes_heading

    def convert_content(
        self, content: str, output_path: str | Path, source_path: Path | None = None
    ) -> list[str]:
        """Отрендерить Markdown и сохранить результат в *output_path*."""
        try:
            document, warnings = GfmDocxRenderer(
                self.default_font_name, self.default_font_size, self.footnotes_heading
            ).render(content, source_path=source_path)
            document.save(str(output_path))
        except Exception as error:
            raise ConversionError(str(error)) from error
        return warnings

    def convert_file(
        self, input_path: str | Path, output_path: str | Path
    ) -> list[str]:
        """Прочитать Markdown-файл и сконвертировать его."""
        source_path = Path(input_path)
        try:
            content = source_path.read_text(encoding="utf-8")
        except OSError as error:
            raise ConversionError(str(error)) from error
        return self.convert_content(content, output_path, source_path)


class WordToMarkdownConverter:
    """Извлекает Markdown из документа Word.

    Преобразование лоссовое: распознаются заголовки по стилям ``Heading N``,
    жирный и курсивный текст по run'ам и таблицы. Списки, изображения,
    формулы и сноски схлопываются в плоский текст.
    """

    def convert_file(
        self, input_path: str | Path, output_path: str | Path
    ) -> list[str]:
        """Сконвертировать документ Word в Markdown-файл."""
        try:
            document = Document(str(input_path))
            lines = self._paragraph_lines(document)
            lines.extend(self._table_lines(document))
            Path(output_path).write_text("\n".join(lines), encoding="utf-8")
        except Exception as error:
            raise ConversionError(str(error)) from error
        return []

    def _paragraph_lines(self, document: Any) -> list[str]:
        lines: list[str] = []
        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                lines.append("")
                continue
            heading_level = self._heading_level(paragraph)
            if heading_level:
                lines.append("#" * heading_level + " " + paragraph.text)
            else:
                lines.append(self._inline_markup(paragraph))
        return lines

    @staticmethod
    def _heading_level(paragraph: Any) -> int:
        """Вернуть уровень заголовка 1..6 или 0, если это не заголовок."""
        style = paragraph.style
        style_name = (style.name or "") if style is not None else ""
        if not style_name.startswith("Heading "):
            return 0
        try:
            level = int(style_name.split()[-1])
        except ValueError:
            return 0
        return level if 1 <= level <= 6 else 0

    @staticmethod
    def _inline_markup(paragraph: Any) -> str:
        parts: list[str] = []
        for run in paragraph.runs:
            text = run.text
            if run.bold and run.italic:
                parts.append(f"***{text}***")
            elif run.bold:
                parts.append(f"**{text}**")
            elif run.italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)
        return "".join(parts)

    @staticmethod
    def _table_lines(document: Any) -> list[str]:
        lines: list[str] = []
        for table in document.tables:
            if not table.rows:
                continue
            lines.append("")
            header = [cell.text.strip() for cell in table.rows[0].cells]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table.rows[1:]:
                lines.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
            lines.append("")
        return lines
