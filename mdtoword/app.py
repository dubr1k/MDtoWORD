import sys
from typing import Any, cast
from pathlib import Path

from docx import Document
from docx.shared import Pt
from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtGui import QDragEnterEvent, QDragMoveEvent, QDropEvent, QIcon, QMouseEvent
from PyQt6.QtWidgets import (
    QApplication, QAbstractItemView, QComboBox, QFileDialog, QGroupBox,
    QHBoxLayout, QLabel, QListWidget, QMainWindow, QMessageBox,
    QPlainTextEdit, QProgressBar, QPushButton, QSpinBox, QTabWidget,
    QVBoxLayout, QWidget,
)

from .workflow import discover_sources, resolve_output_paths
from .gfm_renderer import GfmDocxRenderer
from .theme import ThemeManager


class MarkdownToWordConverter:
    """Convert GFM input to a Word document."""

    def __init__(self):
        self.default_font_name = "Times New Roman"
        self.default_font_size = Pt(12)
        self.footnotes_heading = "Footnotes"

    def convert_content(
        self, content: str, output_path: str | Path, source_path: Path | None = None
    ) -> tuple[bool, str]:
        """Convert Markdown text and save it at *output_path*."""
        try:
            document, warnings = GfmDocxRenderer(
                self.default_font_name, self.default_font_size, self.footnotes_heading
            ).render(content, source_path=source_path)
            document.save(str(output_path))
            message = "Успешно конвертировано"
            if warnings:
                message += "\n\nWarnings:\n" + "\n".join(warnings)
            return True, message
        except Exception as error:
            return False, f"Ошибка при конвертации: {error}"

    def convert_file(
        self, input_path: str | Path, output_path: str | Path
    ) -> tuple[bool, str]:
        """Read and convert a Markdown source file."""
        try:
            source_path = Path(input_path)
            return self.convert_content(
                source_path.read_text(encoding="utf-8"), output_path, source_path
            )
        except Exception as error:
            return False, f"Ошибка при конвертации: {error}"


class WordToMarkdownConverter:
    """Класс для конвертации Word в Markdown"""

    def __init__(self):
        pass

    def convert_file(self, input_path, output_path):
        """Конвертирует Word файл в Markdown"""
        try:
            doc = Document(input_path)

            markdown_lines = []

            for paragraph in doc.paragraphs:
                text = paragraph.text

                # Пропускаем пустые параграфы
                if not text.strip():
                    markdown_lines.append('')
                    continue

                # Определяем уровень заголовка по стилю
                heading_level = 0
                style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
                if style_name.startswith('Heading '):
                    try:
                        heading_level = int(style_name.split()[-1])
                        if 1 <= heading_level <= 6:
                            text = '#' * heading_level + ' ' + text
                    except ValueError:
                        pass  # Если не удалось определить уровень, оставляем как есть

                # Если не заголовок, обрабатываем как обычный текст
                if heading_level == 0:
                    # Простое форматирование текста (жирный, курсив) через run'ы
                    formatted_text = ""
                    for run in paragraph.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            formatted_text += f'***{run_text}***'
                        elif run.bold:
                            formatted_text += f'**{run_text}**'
                        elif run.italic:
                            formatted_text += f'*{run_text}*'
                        elif run_text.strip().startswith('`') and run_text.strip().endswith('`'):
                            formatted_text += f'`{run_text}`'
                        else:
                            formatted_text += run_text

                    text = formatted_text

                markdown_lines.append(text)

            # Обработка таблиц (упрощённо)
            for table in doc.tables:
                if table.rows:
                    markdown_lines.append('')  # Пустая строка перед таблицей
                    # Заголовки (первый ряд)
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
                    # Разделитель
                    markdown_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
                    # Остальные строки
                    for row in table.rows[1:]:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append('| ' + ' | '.join(row_cells) + ' |')
                    markdown_lines.append('')  # Пустая строка после таблицы

            # Записываем результат в файл
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_lines))

            return True, "Успешно конвертировано"

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"


def _dropped_local_paths(event: Any | None) -> list[str]:
    """Local filesystem paths carried by a drop event, if any."""
    if event is None:
        return []
    mime_data = event.mimeData()
    if mime_data is None:
        return []
    return [url.toLocalFile() for url in mime_data.urls() if url.isLocalFile()]


def _accept_local_paths_event(event: Any | None) -> None:
    """Accept a drag event that carries at least one local filesystem path."""
    if event is None:
        return
    if _dropped_local_paths(event):
        event.acceptProposedAction()
    else:
        event.ignore()


class DropFileList(QListWidget):
    """A queue widget that accepts files and directories from the desktop."""

    paths_dropped = pyqtSignal(list)

    def __init__(self, parent: QWidget | None = None):
        super().__init__(parent)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, e: QDragEnterEvent | None) -> None:
        _accept_local_paths_event(e)

    def dragMoveEvent(self, e: QDragMoveEvent | None) -> None:
        _accept_local_paths_event(e)

    def dropEvent(self, event: QDropEvent | None) -> None:
        if event is None:
            return
        paths = _dropped_local_paths(event)
        if paths:
            self.paths_dropped.emit(paths)
            event.acceptProposedAction()
        else:
            event.ignore()


class DropZoneLabel(QLabel):
    """Clickable drop hint that doubles as a file-picker button."""

    clicked = pyqtSignal()

    def __init__(self, parent: QWidget | None = None):
        super().__init__(parent)
        self.setObjectName("drop-zone")
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setWordWrap(True)
        self.setMinimumHeight(80)
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def mouseReleaseEvent(self, event: QMouseEvent | None) -> None:
        if (
            event is not None
            and event.button() == Qt.MouseButton.LeftButton
            and self.rect().contains(event.position().toPoint())
        ):
            self.clicked.emit()
        super().mouseReleaseEvent(event)


class ConverterGUI(QMainWindow):
    """Compact GUI for interactive conversion batches."""

    def __init__(self, theme_manager: ThemeManager | None = None):
        super().__init__()
        self.theme_manager = theme_manager or ThemeManager()
        app = QApplication.instance()
        if isinstance(app, QApplication):
            self.theme_manager.apply(app)
        self.setWindowTitle("MDtoWord")
        self.resize(860, 720)
        self.selected_files: list[Path] = []
        self.output_directory: Path | None = None
        self.current_converter_type = "md_to_word"
        self.converter: MarkdownToWordConverter | WordToMarkdownConverter = MarkdownToWordConverter()
        self.current_language = "ru"
        self.fonts = ["Arial", "Times New Roman", "Calibri", "Georgia", "Helvetica", "Courier New"]
        self.translations = {
            "ru": {
                "title_md": "Markdown → Word", "title_word": "Word → Markdown",
                "settings": "Оформление документа", "font": "Шрифт", "size": "Размер",
                "drop_md": "Перетащите файлы или папки Markdown сюда",
                "drop_word": "Перетащите файлы или папки Word сюда",
                "add_files": "Добавить файлы", "add_folder": "Добавить папку",
                "remove": "Удалить выбранные", "clear": "Очистить очередь",
                "files_tab": "Файлы", "text_tab": "Текст", "text_label": "Введите Markdown-текст",
                "output": "Место сохранения", "output_auto": "Рядом с исходными файлами",
                "choose_output": "Выбрать папку", "reset_output": "Сбросить",
                "ready": "Готово к конвертации", "queued": "В очереди: {count}",
                "converting": "Конвертация: {filename}",
                "finished": "Конвертация завершена", "convert": "Конвертировать",
                "toggle_md": "Режим: MD → Word", "toggle_word": "Режим: Word → MD",
                "theme_dark": "Тёмная тема · Переключить на светлую",
                "theme_light": "Светлая тема · Переключить на тёмную",
                "no_files": "Добавьте файлы или папку для конвертации",
                "empty_text": "Введите текст для конвертации", "save_as": "Сохранить как",
                "errors": "Конвертация завершена с ошибками", "result": "Готово: {success}\nОшибок: {errors}",
                "footnotes_heading": "Сноски",
            },
            "en": {
                "title_md": "Markdown → Word", "title_word": "Word → Markdown",
                "settings": "Document appearance", "font": "Font", "size": "Size",
                "drop_md": "Drop Markdown files or folders here",
                "drop_word": "Drop Word files or folders here",
                "add_files": "Add files", "add_folder": "Add folder",
                "remove": "Remove selected", "clear": "Clear queue",
                "files_tab": "Files", "text_tab": "Text", "text_label": "Enter Markdown text",
                "output": "Save location", "output_auto": "Next to each source file",
                "choose_output": "Choose folder", "reset_output": "Reset",
                "ready": "Ready to convert", "queued": "In queue: {count}",
                "converting": "Converting: {filename}",
                "finished": "Conversion finished", "convert": "Convert",
                "toggle_md": "Mode: MD → Word", "toggle_word": "Mode: Word → MD",
                "theme_dark": "Dark theme · Switch to light",
                "theme_light": "Light theme · Switch to dark",
                "no_files": "Add files or a folder to convert",
                "empty_text": "Enter text to convert", "save_as": "Save as",
                "errors": "Conversion completed with errors", "result": "Complete: {success}\nErrors: {errors}",
                "footnotes_heading": "Footnotes",
            },
        }
        if isinstance(self.converter, MarkdownToWordConverter):
            self.converter.footnotes_heading = self._text["footnotes_heading"]
        self._set_icon()
        self._create_widgets()
        self.setAcceptDrops(True)

    def _set_icon(self) -> None:
        module_dir = Path(__file__).resolve().parent
        # assets/ лежит в корне проекта — на уровень выше пакета mdtoword,
        # а в сборке PyInstaller распаковывается рядом с кодом в sys._MEIPASS.
        search_roots = [module_dir.parent, module_dir]
        bundle_root = getattr(sys, "_MEIPASS", None)
        if bundle_root:
            search_roots.insert(0, Path(bundle_root))
        for icon_name in (("macos-icon.png", "ico.png") if sys.platform == "darwin" else ("ico.png", "macos-icon.png")):
            for root in search_roots:
                icon_path = root / "assets" / icon_name
                if icon_path.exists():
                    self.setWindowIcon(QIcon(str(icon_path)))
                    return

    @property
    def _text(self) -> dict[str, str]:
        return self.translations[self.current_language]

    def dragEnterEvent(self, event: QDragEnterEvent | None) -> None:
        _accept_local_paths_event(event)

    def dragMoveEvent(self, event: QDragMoveEvent | None) -> None:
        _accept_local_paths_event(event)

    def dropEvent(self, event: QDropEvent | None) -> None:
        if event is None:
            return
        paths = _dropped_local_paths(event)
        if paths:
            self._add_sources(paths)
            event.acceptProposedAction()
        else:
            event.ignore()

    def _create_widgets(self) -> None:
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(20, 14, 20, 16)
        layout.setSpacing(10)

        self.title_label = QLabel()
        self.title_label.setObjectName("title-label")
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.title_label)

        self.settings_group = QGroupBox()
        settings = QHBoxLayout(self.settings_group)
        self.font_label = QLabel()
        self.font_combobox = QComboBox()
        self.font_combobox.addItems(self.fonts)
        self.font_combobox.setCurrentText("Times New Roman")
        self.font_combobox.currentTextChanged.connect(self._on_font_change)
        self.size_label = QLabel()
        self.size_spinbox = QSpinBox()
        self.size_spinbox.setRange(6, 72)
        self.size_spinbox.setValue(12)
        self.size_spinbox.valueChanged.connect(self._on_size_change)
        settings.addWidget(self.font_label)
        settings.addWidget(self.font_combobox, 1)
        settings.addWidget(self.size_label)
        settings.addWidget(self.size_spinbox)
        layout.addWidget(self.settings_group)

        self.tabs = QTabWidget()
        layout.addWidget(self.tabs, 1)
        self.files_tab = QWidget()
        self.files_tab.setObjectName("tab-page")
        files_layout = QVBoxLayout(self.files_tab)
        files_layout.setContentsMargins(16, 16, 16, 16)
        files_layout.setSpacing(10)
        self.drop_hint = DropZoneLabel()
        self.drop_hint.clicked.connect(self._select_files)
        files_layout.addWidget(self.drop_hint)
        actions = QHBoxLayout()
        self.add_files_button = QPushButton()
        self.add_files_button.clicked.connect(self._select_files)
        self.add_folder_button = QPushButton()
        self.add_folder_button.clicked.connect(self._select_folder)
        actions.addWidget(self.add_files_button)
        actions.addWidget(self.add_folder_button)
        actions.addStretch()
        files_layout.addLayout(actions)
        self.files_listbox = DropFileList()
        self.files_listbox.paths_dropped.connect(self._add_sources)
        self.files_listbox.itemSelectionChanged.connect(self._update_queue_buttons)
        self.files_listbox.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.files_listbox.setMinimumHeight(90)
        files_layout.addWidget(self.files_listbox, 1)
        removal = QHBoxLayout()
        self.remove_button = QPushButton()
        self.remove_button.setObjectName("danger-button")
        self.remove_button.clicked.connect(self._remove_selected_files)
        self.clear_button = QPushButton()
        self.clear_button.setObjectName("danger-button")
        self.clear_button.clicked.connect(self._clear_files)
        removal.addWidget(self.remove_button)
        removal.addWidget(self.clear_button)
        removal.addStretch()
        files_layout.addLayout(removal)
        self.tabs.addTab(self.files_tab, "")

        self.text_tab = QWidget()
        self.text_tab.setObjectName("tab-page")
        text_layout = QVBoxLayout(self.text_tab)
        self.text_label = QLabel()
        self.text_input = QPlainTextEdit()
        text_layout.addWidget(self.text_label)
        text_layout.addWidget(self.text_input)
        self.tabs.addTab(self.text_tab, "")

        self.output_group = QGroupBox()
        output = QHBoxLayout(self.output_group)
        self.output_label = QLabel()
        self.output_label.setObjectName("output-path")
        self.choose_output_button = QPushButton()
        self.choose_output_button.clicked.connect(self._select_output_directory)
        self.reset_output_button = QPushButton()
        self.reset_output_button.clicked.connect(self._reset_output_directory)
        output.addWidget(self.output_label, 1)
        output.addWidget(self.choose_output_button)
        output.addWidget(self.reset_output_button)
        layout.addWidget(self.output_group)

        self.progress = QProgressBar()
        policy = self.progress.sizePolicy()
        policy.setRetainSizeWhenHidden(True)
        self.progress.setSizePolicy(policy)
        self.progress.hide()
        layout.addWidget(self.progress)
        self.status_label = QLabel()
        self.status_label.setObjectName("status-label")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.status_label)
        self.convert_button = QPushButton()
        self.convert_button.setObjectName("primary-button")
        self.convert_button.clicked.connect(self._convert_files)
        layout.addWidget(self.convert_button)

        footer = QHBoxLayout()
        self.toggle_button = QPushButton()
        self.toggle_button.clicked.connect(self._toggle_converter_type)
        self.language_button = QPushButton("EN")
        self.language_button.clicked.connect(self._toggle_language)
        self.theme_button = QPushButton()
        self.theme_button.setObjectName("theme-button")
        self.theme_button.clicked.connect(self._toggle_theme)
        footer.addWidget(self.toggle_button)
        footer.addStretch()
        footer.addWidget(self.theme_button)
        footer.addWidget(self.language_button)
        self._update_theme_button()
        layout.addLayout(footer)
        self._update_ui()

    def _toggle_converter_type(self) -> None:
        self.current_converter_type = (
            "word_to_md" if self.current_converter_type == "md_to_word" else "md_to_word"
        )
        self.converter = (
            WordToMarkdownConverter() if self.current_converter_type == "word_to_md" else MarkdownToWordConverter()
        )
        if isinstance(self.converter, MarkdownToWordConverter):
            self.converter.default_font_name = self.font_combobox.currentText()
            self.converter.default_font_size = Pt(self.size_spinbox.value())
            self.converter.footnotes_heading = self._text["footnotes_heading"]
        self.selected_files = discover_sources(self.selected_files, self.current_converter_type)
        self._update_ui()

    def _toggle_language(self) -> None:
        self.current_language = "en" if self.current_language == "ru" else "ru"
        if isinstance(self.converter, MarkdownToWordConverter):
            self.converter.footnotes_heading = self._text["footnotes_heading"]
        self._update_ui()

    def _toggle_theme(self) -> None:
        self.theme_manager.toggle()
        app = QApplication.instance()
        if isinstance(app, QApplication):
            self.theme_manager.apply(app)
        self._update_theme_button()

    def _update_theme_button(self) -> None:
        is_dark = self.theme_manager.theme == "dark"
        self.theme_button.setText("☀" if is_dark else "☾")
        tooltip = self._text["theme_dark" if is_dark else "theme_light"]
        self.theme_button.setToolTip(tooltip)
        self.theme_button.setAccessibleName(tooltip)

    def _update_ui(self) -> None:
        text = self._text
        is_markdown = self.current_converter_type == "md_to_word"
        self.setWindowTitle(text["title_md"] if is_markdown else text["title_word"])
        self.title_label.setText(self.windowTitle())
        self.settings_group.setTitle(text["settings"])
        self.settings_group.setVisible(is_markdown)
        self.font_label.setText(text["font"])
        self.size_label.setText(text["size"])
        self.drop_hint.setText(text["drop_md"] if is_markdown else text["drop_word"])
        self.add_files_button.setText(text["add_files"])
        self.add_folder_button.setText(text["add_folder"])
        self.remove_button.setText(text["remove"])
        self.clear_button.setText(text["clear"])
        self.tabs.setTabText(self.tabs.indexOf(self.files_tab), text["files_tab"])
        self.tabs.setTabText(self.tabs.indexOf(self.text_tab), text["text_tab"])
        self.text_label.setText(text["text_label"])
        self.tabs.setTabVisible(self.tabs.indexOf(self.text_tab), is_markdown)
        self.output_group.setTitle(text["output"])
        self.output_label.setText(str(self.output_directory) if self.output_directory else text["output_auto"])
        self.choose_output_button.setText(text["choose_output"])
        self.reset_output_button.setText(text["reset_output"])
        self.reset_output_button.setVisible(self.output_directory is not None)
        self.toggle_button.setText(text["toggle_md"] if is_markdown else text["toggle_word"])
        self.language_button.setText("EN" if self.current_language == "ru" else "RU")
        self._update_theme_button()
        self._refresh_queue()

    def _refresh_queue(self) -> None:
        self.files_listbox.clear()
        for source in self.selected_files:
            self.files_listbox.addItem(f"{source.name}\n{source.parent}")
        count = len(self.selected_files)
        if count:
            self.status_label.setText(self._text["queued"].format(count=count))
        else:
            self.status_label.setText(self._text["ready"])
        self.convert_button.setText(
            self._text["convert"] if not count else f"{self._text['convert']} ({count})"
        )
        self._update_queue_buttons()

    def _update_queue_buttons(self) -> None:
        self.clear_button.setEnabled(bool(self.selected_files))
        self.remove_button.setEnabled(bool(self.files_listbox.selectedItems()))

    def _on_font_change(self, font_name: str) -> None:
        if isinstance(self.converter, MarkdownToWordConverter):
            self.converter.default_font_name = font_name

    def _on_size_change(self, value: int) -> None:
        if isinstance(self.converter, MarkdownToWordConverter):
            self.converter.default_font_size = Pt(value)

    def _select_files(self) -> None:
        suffix = "*.md *.markdown" if self.current_converter_type == "md_to_word" else "*.docx"
        paths, _ = QFileDialog.getOpenFileNames(self, self._text["add_files"], "", f"Supported files ({suffix})")
        self._add_sources(paths)

    def _select_folder(self) -> None:
        folder = QFileDialog.getExistingDirectory(self, self._text["add_folder"])
        if folder:
            self._add_sources([folder])

    def _add_sources(self, paths: list[str] | list[Path]) -> None:
        discovered = discover_sources((Path(path) for path in paths), self.current_converter_type)
        known = set(self.selected_files)
        self.selected_files.extend(path for path in discovered if path not in known)
        self._refresh_queue()

    def _remove_selected_files(self) -> None:
        selected = {index.row() for index in self.files_listbox.selectedIndexes()}
        self.selected_files = [
            source for index, source in enumerate(self.selected_files) if index not in selected
        ]
        self._refresh_queue()

    def _clear_files(self) -> None:
        self.selected_files.clear()
        self._refresh_queue()

    def _select_output_directory(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, self._text["choose_output"])
        if directory:
            self.output_directory = Path(directory)
            self._update_ui()

    def _reset_output_directory(self) -> None:
        self.output_directory = None
        self._update_ui()

    def _convert_files(self) -> None:
        text_tab_index = self.tabs.indexOf(self.text_tab)
        if self.current_converter_type == "md_to_word" and self.tabs.currentIndex() == text_tab_index:
            self._convert_text()
            return
        if not self.selected_files:
            QMessageBox.warning(self, self.windowTitle(), self._text["no_files"])
            return

        suffix = ".docx" if self.current_converter_type == "md_to_word" else ".md"
        queue = list(self.selected_files)
        outputs = resolve_output_paths(queue, self.output_directory, suffix)

        lockable_widgets = (
            self.convert_button,
            self.files_listbox,
            self.add_files_button,
            self.add_folder_button,
            self.remove_button,
            self.clear_button,
            self.drop_hint,
        )
        try:
            for widget in lockable_widgets:
                widget.setEnabled(False)
            self.setAcceptDrops(False)
            self.files_listbox.setAcceptDrops(False)
            self.progress.show()
            self.progress.setRange(0, len(queue))
            success_count = 0
            errors: list[str] = []
            warnings: list[str] = []
            for index, source in enumerate(queue, start=1):
                self.status_label.setText(self._text["converting"].format(filename=source.name))
                QApplication.processEvents()
                success, message = self.converter.convert_file(source, outputs[source])
                if success:
                    success_count += 1
                    if "Warnings:" in message:
                        warnings.append(f"{source.name}: {message.split('Warnings:', 1)[1].strip()}")
                else:
                    errors.append(f"{source.name}: {message}")
                self.progress.setValue(index)
                QApplication.processEvents()

            details = errors + warnings
            result = self._text["result"].format(success=success_count, errors=len(errors))
            if details:
                QMessageBox.warning(self, self._text["errors"], result + "\n\n" + "\n".join(details))
            else:
                QMessageBox.information(self, self.windowTitle(), result)
            self.status_label.setText(self._text["finished"])
        finally:
            self.progress.hide()
            self.progress.reset()
            self.setAcceptDrops(True)
            self.files_listbox.setAcceptDrops(True)
            for widget in lockable_widgets:
                widget.setEnabled(True)
            self._update_queue_buttons()

    def _convert_text(self) -> None:
        if not isinstance(self.converter, MarkdownToWordConverter):
            return
        content = self.text_input.toPlainText()
        if not content.strip():
            QMessageBox.warning(self, self.windowTitle(), self._text["empty_text"])
            return
        output_path, _ = QFileDialog.getSaveFileName(self, self._text["save_as"], "", "Word files (*.docx)")
        if not output_path:
            return
        output = Path(output_path).with_suffix(".docx")
        success, message = cast(MarkdownToWordConverter, self.converter).convert_content(content, output)
        if success:
            QMessageBox.information(self, self.windowTitle(), message)
        else:
            QMessageBox.critical(self, self._text["errors"], message)


def main():
    """Главная функция"""
    app = QApplication(sys.argv)
    window = ConverterGUI()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
