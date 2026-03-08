import sys
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QComboBox, QSpinBox, QTabWidget,
    QGroupBox, QListWidget, QProgressBar, QPlainTextEdit,
    QFileDialog, QMessageBox, QAbstractItemView
)
from PyQt6.QtGui import QIcon, QFont
from PyQt6.QtCore import Qt


class MarkdownToWordConverter:
    """Класс для конвертации Markdown в Word"""

    def __init__(self):
        self.doc = None
        self.default_font_name = 'Times New Roman'
        self.default_font_size = Pt(12)

    def create_document(self):
        """Создает новый документ Word"""
        self.doc = Document()
        # Устанавливаем стандартный стиль
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.default_font_name
        font.size = self.default_font_size
        font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

    def add_table_borders(self, table):
        """Добавляет границы к таблице"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Создаем элемент границ
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def parse_inline_formatting(self, text):
        """Парсит встроенное форматирование (жирный, курсив, код)"""
        # Возвращает список кортежей (текст, форматирование)
        # форматирование: {'bold': bool, 'italic': bool, 'code': bool}
        parts = []

        # Регулярное выражение для поиска форматирования
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|___.*?___|__.*?__|_.*?_)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Добавляем текст до совпадения
            if match.start() > last_end:
                parts.append((text[last_end:match.start()], {}))

            matched_text = match.group()
            formatting = {}
            clean_text = matched_text

            # Проверяем тип форматирования
            if matched_text.startswith('***') or matched_text.startswith('___'):
                formatting = {'bold': True, 'italic': True}
                clean_text = matched_text[3:-3]
            elif matched_text.startswith('**') or matched_text.startswith('__'):
                formatting = {'bold': True}
                clean_text = matched_text[2:-2]
            elif matched_text.startswith('*') or matched_text.startswith('_'):
                formatting = {'italic': True}
                clean_text = matched_text[1:-1]
            elif matched_text.startswith('`'):
                formatting = {'code': True}
                clean_text = matched_text[1:-1]

            parts.append((clean_text, formatting))
            last_end = match.end()

        # Добавляем оставшийся текст
        if last_end < len(text):
            parts.append((text[last_end:], {}))

        return parts if parts else [(text, {})]

    def add_formatted_text(self, paragraph, text):
        """Добавляет текст с форматированием в параграф"""
        parts = self.parse_inline_formatting(text)

        for part_text, formatting in parts:
            run = paragraph.add_run(part_text)
            run.font.name = self.default_font_name
            run.font.size = self.default_font_size
            run.font.color.rgb = RGBColor(0, 0, 0)

            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True
            if formatting.get('code'):
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    def process_table(self, lines, start_idx):
        """Обрабатывает таблицу из markdown"""
        table_lines = []
        idx = start_idx

        # Собираем все строки таблицы
        while idx < len(lines) and '|' in lines[idx]:
            table_lines.append(lines[idx])
            idx += 1

        if len(table_lines) < 2:
            return idx

        # Парсим таблицу
        rows = []
        for line in table_lines:
            # Пропускаем разделительную строку (---|---|---)
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue

            # Разбиваем строку на ячейки
            cells = [cell.strip() for cell in line.split('|')]
            # Удаляем пустые ячейки в начале и конце
            cells = [c for c in cells if c or cells.index(c) not in [0, len(cells) - 1]]
            if cells:
                rows.append(cells)

        if not rows:
            return idx

        # Создаем таблицу в документе
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        self.add_table_borders(table)

        # Заполняем таблицу
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    # Очищаем ячейку и добавляем форматированный текст
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    self.add_formatted_text(paragraph, cell_text)
                    # Устанавливаем выравнивание по ширине для текста в ячейках
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # Форматирование для заголовка таблицы (первая строка)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

        return idx

    def process_list(self, lines, start_idx):
        """Обрабатывает списки (маркированные и нумерованные)"""
        idx = start_idx
        list_items = []

        # Определяем тип списка
        first_line = lines[start_idx].strip()
        is_ordered = bool(re.match(r'^\d+\.', first_line))

        # Собираем элементы списка
        while idx < len(lines):
            line = lines[idx].strip()
            if not line:
                break

            # Проверяем маркированный список
            if re.match(r'^[-*+]\s', line):
                list_items.append(('unordered', line[2:]))
                idx += 1
            # Проверяем нумерованный список
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^\d+\.\s(.*)', line)
                list_items.append(('ordered', match.group(1)))
                idx += 1
            else:
                break

        # Добавляем элементы списка в документ
        for list_type, item_text in list_items:
            paragraph = self.doc.add_paragraph(
                style='List Bullet' if list_type == 'unordered' else 'List Number'
            )
            self.add_formatted_text(paragraph, item_text)

        return idx

    def convert_content(self, content, output_path):
        """Конвертирует строку markdown в Word"""
        try:
            # Создаем документ
            self.create_document()

            # Разбиваем на строки
            lines = content.split('\n')

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # Пропускаем пустые строки
                if not stripped:
                    i += 1
                    continue

                # Заголовки
                if stripped.startswith('#'):
                    level = 0
                    while level < len(stripped) and stripped[level] == '#':
                        level += 1

                    title_text = stripped[level:].strip()
                    heading_style = f'Heading {min(level, 9)}'

                    paragraph = self.doc.add_paragraph(style=heading_style)
                    self.add_formatted_text(paragraph, title_text)
                    paragraph.paragraph_format.space_after = Pt(6)
                    i += 1

                # Горизонтальная линия
                elif stripped in ('---', '***', '___'):
                    self.doc.add_paragraph('_' * 50)
                    i += 1

                # Таблицы
                elif '|' in line:
                    i = self.process_table(lines, i)

                # Списки
                elif re.match(r'^[-*+]\s', stripped) or re.match(r'^\d+\.\s', stripped):
                    i = self.process_list(lines, i)

                # Блок кода (```)
                elif stripped.startswith('```'):
                    i += 1
                    code_lines = []
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1

                    if code_lines:
                        paragraph = self.doc.add_paragraph()
                        run = paragraph.add_run('\n'.join(code_lines))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    i += 1

                # Обычный текст
                else:
                    paragraph = self.doc.add_paragraph()
                    self.add_formatted_text(paragraph, line)
                    # Устанавливаем выравнивание по ширине для обычного текста
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    i += 1

            # Сохраняем документ
            self.doc.save(output_path)
            return True, "Успешно конвертировано"

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"

    def convert_file(self, input_path, output_path):
        """Конвертирует markdown файл в Word"""
        try:
            # Читаем файл
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return self.convert_content(content, output_path)

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"


class WordToMarkdownConverter:
    """Класс для конвертации Word в Markdown"""

    def __init__(self):
        pass

    def convert_file(self, input_path, output_path):
        """Конвертирует Word файл в Markdown"""
        try:
            doc = Document(input_path)

            markdown_lines = []

            for paragraph in doc.paragraphs:
                text = paragraph.text

                # Пропускаем пустые параграфы
                if not text.strip():
                    markdown_lines.append('')
                    continue

                # Определяем уровень заголовка по стилю
                heading_level = 0
                if paragraph.style.name.startswith('Heading '):
                    try:
                        heading_level = int(paragraph.style.name.split()[-1])
                        if 1 <= heading_level <= 6:
                            text = '#' * heading_level + ' ' + text
                    except ValueError:
                        pass  # Если не удалось определить уровень, оставляем как есть

                # Если не заголовок, обрабатываем как обычный текст
                if heading_level == 0:
                    # Простое форматирование текста (жирный, курсив) через run'ы
                    formatted_text = ""
                    for run in paragraph.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            formatted_text += f'***{run_text}***'
                        elif run.bold:
                            formatted_text += f'**{run_text}**'
                        elif run.italic:
                            formatted_text += f'*{run_text}*'
                        elif run_text.strip().startswith('`') and run_text.strip().endswith('`'):
                            formatted_text += f'`{run_text}`'
                        else:
                            formatted_text += run_text

                    text = formatted_text

                markdown_lines.append(text)

            # Обработка таблиц (упрощённо)
            for table in doc.tables:
                if table.rows:
                    markdown_lines.append('')  # Пустая строка перед таблицей
                    # Заголовки (первый ряд)
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_lines.append('| ' + ' | '.join(header_cells) + ' |')
                    # Разделитель
                    markdown_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
                    # Остальные строки
                    for row in table.rows[1:]:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append('| ' + ' | '.join(row_cells) + ' |')
                    markdown_lines.append('')  # Пустая строка после таблицы

            # Записываем результат в файл
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_lines))

            return True, "Успешно конвертировано"

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"


class ConverterGUI(QMainWindow):
    """GUI для конвертера на PyQt6"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Конвертер Markdown в Word")
        self.resize(800, 700)

        # Устанавливаем иконку приложения (assets/ в корне проекта)
        script_dir = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(script_dir, 'assets', 'ico.png')
        if not os.path.exists(icon_path):
            icon_path = os.path.join(script_dir, 'ico.png')  # fallback: иконка в корне
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        self.selected_files = []
        self.output_directory = ""
        self.converter = MarkdownToWordConverter()
        self.current_converter_type = "md_to_word"

        # Популярные шрифты (для режима MD->Word)
        self.fonts = [
            "Arial", "Times New Roman", "Courier New", "Helvetica", "Verdana",
            "Georgia", "Palatino", "Garamond", "Comic Sans MS", "Trebuchet MS",
            "Arial Black", "Impact", "Calibri", "Cambria", "Candara",
            "Consolas", "Segoe UI", "Roboto", "Open Sans", "Lato",
            "Montserrat", "Source Sans Pro", "Noto Sans", "Ubuntu", "Merriweather",
            "Lora", "Poppins", "Inter", "Fira Sans", "PT Sans",
            "Droid Sans", "Nunito", "Raleway", "Oswald", "Roboto Condensed",
            "Titillium Web", "Exo 2", "Crimson Text", "Playfair Display", "Quicksand",
            "Comfortaa", "Pacifico", "Indie Flower", "Dancing Script", "Kaushan Script",
            "Courgette", "Satisfy", "Handlee", "Bangers", "Chewy",
            "Fredoka One", "Architects Daughter", "Nova Square", "Orbitron", "Press Start 2P",
            "Monoton", "VT323", "Cutive Mono", "Inconsolata", "Space Mono",
            "Fira Mono", "Source Code Pro", "Roboto Mono", "IBM Plex Mono", "JetBrains Mono",
            "Cascadia Code", "Ubuntu Mono", "PT Mono", "Noto Sans Mono", "Roboto Slab",
            "Crimson Pro", "Barlow", "Nunito Sans", "Work Sans", "Public Sans",
            "Red Hat Display", "Red Hat Text", "Atkinson Hyperlegible", "Literata", "Charter",
            "Avenir Next", "Proxima Nova", "Myriad Pro", "Gill Sans", "Franklin Gothic",
            "Optima", "Futura", "Univers", "Helvetica Neue", "Avenir",
            "Geneva", "Tahoma", "Trebuchet MS", "Lucida Grande", "Lucida Sans Unicode",
            "Bitstream Vera Sans", "DejaVu Sans", "Liberation Sans", "Noto Sans CJK",
            "Source Han Sans", "Noto Serif CJK", "Source Han Serif", "SimSun", "SimHei",
            "Microsoft YaHei", "Microsoft JhengHei", "Meiryo", "Yu Gothic",
            "Hiragino Kaku Gothic Pro", "Apple Gothic", "Malgun Gothic", "Batang", "Dotum",
            "Arial Unicode MS", "Lucida Sans", "Segoe UI Symbol", "Symbol", "Webdings",
            "Wingdings", "Wingdings 2", "Wingdings 3",
        ]

        # Словарь переводов
        self.translations = {
            "ru": {
                "title": "Конвертер",
                "title_md_to_word": "Конвертер Markdown в Word",
                "title_word_to_md": "Конвертер Word в Markdown",
                "settings_frame": "Настройки документа",
                "font_label": "Шрифт:",
                "width_label": "Ширина текста (пт):",
                "files_frame": "Выбор файлов",
                "files_frame_md": "Выбор файлов Markdown",
                "files_frame_word": "Выбор файлов Word",
                "select_button": "Выбрать файлы .md",
                "select_button_md": "Выбрать файлы .md",
                "select_button_word": "Выбрать файлы .docx",
                "remove_button": "Удалить выбранные",
                "remove_all_button": "Удалить все",
                "tab_files": "Файлы",
                "tab_text": "Текст",
                "text_input_label": "Введите Markdown текст:",
                "save_as_title": "Сохранить как",
                "output_frame": "Место сохранения",
                "output_label_default": "Папка не выбрана",
                "output_button": "Выбрать папку",
                "status_ready": "Готов к конвертации",
                "convert_button": "Конвертировать",
                "toggle_button_md_to_word": "Режим: MD -> Word",
                "toggle_button_word_to_md": "Режим: Word -> MD",
                "warning_no_files": "Выберите файлы для конвертации",
                "warning_no_dir": "Выберите папку для сохранения",
                "success_message": "Все файлы успешно конвертированы!\nВсего: {count}\nШрифт: {font}, Размер: {size} pt",
                "success_message_word": "Все файлы успешно конвертированы!\nВсего: {count}",
                "status_converting": "Конвертация: {filename}",
                "status_finished": "Конвертация завершена",
                "error_title": "Конвертация завершена с ошибками",
                "error_message": "Успешно: {success}\nОшибок: {errors}\n\n{details}",
                "font_changed": "Шрифт изменен на: {font}",
                "size_changed": "Размер шрифта изменен на: {size} pt",
                "invalid_size": "Некорректное значение ширины (размера шрифта)",
                "files_selected": "Выбрано файлов: {count}",
                "icon_failed": "Не удалось установить иконку: {error}",
                "logo_failed": "Не удалось загрузить логотип: {error}",
            },
            "en": {
                "title": "Converter",
                "title_md_to_word": "Markdown to Word Converter",
                "title_word_to_md": "Word to Markdown Converter",
                "settings_frame": "Document Settings",
                "font_label": "Font:",
                "width_label": "Text Width (pt):",
                "files_frame": "Select Files",
                "files_frame_md": "Select Markdown Files",
                "files_frame_word": "Select Word Files",
                "select_button": "Select .md Files",
                "select_button_md": "Select .md Files",
                "select_button_word": "Select .docx Files",
                "remove_button": "Remove Selected",
                "remove_all_button": "Remove All",
                "tab_files": "Files",
                "tab_text": "Text",
                "text_input_label": "Enter Markdown Text:",
                "save_as_title": "Save As",
                "output_frame": "Save Location",
                "output_label_default": "Folder not selected",
                "output_button": "Select Folder",
                "status_ready": "Ready to convert",
                "convert_button": "Convert",
                "toggle_button_md_to_word": "Mode: MD -> Word",
                "toggle_button_word_to_md": "Mode: Word -> MD",
                "warning_no_files": "Select files for conversion",
                "warning_no_dir": "Select a folder to save",
                "success_message": "All files converted successfully!\nTotal: {count}\nFont: {font}, Size: {size} pt",
                "success_message_word": "All files converted successfully!\nTotal: {count}",
                "status_converting": "Converting: {filename}",
                "status_finished": "Conversion finished",
                "error_title": "Conversion completed with errors",
                "error_message": "Success: {success}\nErrors: {errors}\n\n{details}",
                "font_changed": "Font changed to: {font}",
                "size_changed": "Font size changed to: {size} pt",
                "invalid_size": "Invalid width (font size) value",
                "files_selected": "Files selected: {count}",
                "icon_failed": "Failed to set icon: {error}",
                "logo_failed": "Failed to load logo: {error}",
            }
        }
        self.current_language = "ru"

        self._create_widgets()

    def _create_widgets(self):
        """Создает виджеты интерфейса"""
        t = self.translations[self.current_language]

        # Центральный виджет и основной лейаут
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        # --- Заголовок ---
        self.title_label = QLabel(t["title_md_to_word"])
        self.title_label.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        self.title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(self.title_label)

        # --- Настройки документа (QGroupBox) ---
        self.settings_group = QGroupBox(t["settings_frame"])
        settings_layout = QHBoxLayout(self.settings_group)

        self.font_label_widget = QLabel(t["font_label"])
        settings_layout.addWidget(self.font_label_widget)

        self.font_combobox = QComboBox()
        self.font_combobox.addItems(self.fonts)
        self.font_combobox.setCurrentText(self.converter.default_font_name)
        self.font_combobox.setMinimumWidth(180)
        self.font_combobox.currentTextChanged.connect(self._on_font_change)
        settings_layout.addWidget(self.font_combobox)

        self.width_label_widget = QLabel(t["width_label"])
        settings_layout.addWidget(self.width_label_widget)

        self.width_spinbox = QSpinBox()
        self.width_spinbox.setRange(1, 100)
        self.width_spinbox.setValue(int(self.converter.default_font_size.pt))
        self.width_spinbox.valueChanged.connect(self._on_width_change)
        settings_layout.addWidget(self.width_spinbox)

        settings_layout.addStretch()
        main_layout.addWidget(self.settings_group)

        # --- Вкладки ---
        self.notebook = QTabWidget()
        main_layout.addWidget(self.notebook)

        # -- Вкладка «Файлы» --
        tab_files_widget = QWidget()
        tab_files_layout = QVBoxLayout(tab_files_widget)

        self.files_group = QGroupBox(t["files_frame_md"])
        files_group_layout = QVBoxLayout(self.files_group)

        # Кнопка выбора файлов
        self.select_button = QPushButton(t["select_button_md"])
        self.select_button.setStyleSheet(
            "QPushButton { background-color: #4CAF50; color: white; "
            "font: bold 10pt 'Arial'; padding: 5px 10px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #45a049; }"
        )
        self.select_button.clicked.connect(self._select_files)
        files_group_layout.addWidget(self.select_button)

        # Список файлов
        self.files_listbox = QListWidget()
        self.files_listbox.setSelectionMode(
            QAbstractItemView.SelectionMode.ExtendedSelection
        )
        self.files_listbox.setFont(QFont("Arial", 9))
        self.files_listbox.setMinimumHeight(150)
        files_group_layout.addWidget(self.files_listbox)

        # Кнопки удаления
        remove_layout = QHBoxLayout()

        self.remove_button = QPushButton(t["remove_button"])
        self.remove_button.setStyleSheet(
            "QPushButton { background-color: #f44336; color: white; "
            "font: 9pt 'Arial'; padding: 3px 5px; border-radius: 3px; }"
            "QPushButton:hover { background-color: #e53935; }"
        )
        self.remove_button.clicked.connect(self._remove_selected_files)
        remove_layout.addWidget(self.remove_button)

        self.remove_all_button = QPushButton(t["remove_all_button"])
        self.remove_all_button.setStyleSheet(
            "QPushButton { background-color: #d32f2f; color: white; "
            "font: 9pt 'Arial'; padding: 3px 5px; border-radius: 3px; }"
            "QPushButton:hover { background-color: #c62828; }"
        )
        self.remove_all_button.clicked.connect(self._remove_all_files)
        remove_layout.addWidget(self.remove_all_button)

        remove_layout.addStretch()
        files_group_layout.addLayout(remove_layout)
        tab_files_layout.addWidget(self.files_group)

        self.notebook.addTab(tab_files_widget, t["tab_files"])

        # -- Вкладка «Текст» (только для режима MD->Word) --
        self.tab_text_widget = QWidget()
        tab_text_layout = QVBoxLayout(self.tab_text_widget)

        self.text_input_label = QLabel(t["text_input_label"])
        tab_text_layout.addWidget(self.text_input_label)

        self.text_input = QPlainTextEdit()
        self.text_input.setFont(QFont("Arial", 10))
        tab_text_layout.addWidget(self.text_input)

        self.notebook.addTab(self.tab_text_widget, t["tab_text"])

        # --- Место сохранения ---
        self.output_group = QGroupBox(t["output_frame"])
        output_layout = QHBoxLayout(self.output_group)

        self.output_label = QLabel(t["output_label_default"])
        self.output_label.setStyleSheet(
            "background-color: #555; color: white; padding: 5px; border-radius: 3px;"
        )
        output_layout.addWidget(self.output_label, stretch=1)

        self.output_button = QPushButton(t["output_button"])
        self.output_button.setStyleSheet(
            "QPushButton { background-color: #2196F3; color: white; "
            "font: bold 10pt 'Arial'; padding: 5px 10px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #1e88e5; }"
        )
        self.output_button.clicked.connect(self._select_output_directory)
        output_layout.addWidget(self.output_button)

        main_layout.addWidget(self.output_group)

        # --- Прогресс-бар ---
        self.progress = QProgressBar()
        self.progress.setValue(0)
        main_layout.addWidget(self.progress)

        # --- Статус ---
        self.status_label = QLabel(t["status_ready"])
        self.status_label.setFont(QFont("Arial", 9))
        self.status_label.setStyleSheet("color: gray;")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(self.status_label)

        # --- Кнопка конвертации ---
        self.convert_button = QPushButton(t["convert_button"])
        self.convert_button.setStyleSheet(
            "QPushButton { background-color: #FF9800; color: white; "
            "font: bold 12pt 'Arial'; padding: 10px 20px; border-radius: 5px; }"
            "QPushButton:hover { background-color: #fb8c00; }"
        )
        self.convert_button.clicked.connect(self._convert_files)
        main_layout.addWidget(self.convert_button)

        # --- Нижняя панель ---
        bottom_layout = QHBoxLayout()

        self.toggle_button = QPushButton(t["toggle_button_md_to_word"])
        self.toggle_button.setStyleSheet(
            "QPushButton { background-color: #670067; color: white; "
            "font: 12pt 'Arial'; padding: 3px 10px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #7b007b; }"
        )
        self.toggle_button.clicked.connect(self._toggle_converter_type)
        bottom_layout.addWidget(self.toggle_button)

        bottom_layout.addStretch()

        self.language_button = QPushButton("EN")
        self.language_button.setStyleSheet(
            "QPushButton { background-color: #9E9E9E; color: white; "
            "font: 12pt 'Arial'; padding: 3px 10px; border-radius: 4px; }"
            "QPushButton:hover { background-color: #8e8e8e; }"
        )
        self.language_button.clicked.connect(self._toggle_language)
        bottom_layout.addWidget(self.language_button)

        main_layout.addLayout(bottom_layout)

        # Инициализируем видимость элементов
        self._update_ui_for_converter_type()

    # ------------------------------------------------------------------ #
    #  Переключение режимов и языка
    # ------------------------------------------------------------------ #

    def _toggle_converter_type(self):
        """Переключает между режимами MD->Word и Word->MD"""
        if self.current_converter_type == "md_to_word":
            self.current_converter_type = "word_to_md"
            self.converter = WordToMarkdownConverter()
        else:
            self.current_converter_type = "md_to_word"
            self.converter = MarkdownToWordConverter()
        self._update_ui_for_converter_type()

    def _update_ui_for_converter_type(self):
        """Обновляет интерфейс в зависимости от режима конвертации"""
        t = self.translations[self.current_language]
        if self.current_converter_type == "md_to_word":
            self.setWindowTitle(t["title_md_to_word"])
            self.title_label.setText(t["title_md_to_word"])
            self.toggle_button.setText(t["toggle_button_md_to_word"])
            self.files_group.setTitle(t["files_frame_md"])
            self.select_button.setText(t["select_button_md"])
            self.settings_group.setVisible(True)
            # Показываем вкладку «Текст», если её нет
            if self.notebook.indexOf(self.tab_text_widget) == -1:
                self.notebook.addTab(self.tab_text_widget, t["tab_text"])
        else:
            self.setWindowTitle(t["title_word_to_md"])
            self.title_label.setText(t["title_word_to_md"])
            self.toggle_button.setText(t["toggle_button_word_to_md"])
            self.files_group.setTitle(t["files_frame_word"])
            self.select_button.setText(t["select_button_word"])
            self.settings_group.setVisible(False)
            # Убираем вкладку «Текст»
            text_tab_idx = self.notebook.indexOf(self.tab_text_widget)
            if text_tab_idx != -1:
                self.notebook.removeTab(text_tab_idx)

    def _toggle_language(self):
        """Переключает язык интерфейса"""
        if self.current_language == "ru":
            self.current_language = "en"
            self.language_button.setText("RU")
        else:
            self.current_language = "ru"
            self.language_button.setText("EN")
        self._update_ui_text()

    def _update_ui_text(self):
        """Обновляет текст всех виджетов под текущий язык"""
        t = self.translations[self.current_language]

        # Заголовки / режим
        if self.current_converter_type == "md_to_word":
            self.setWindowTitle(t["title_md_to_word"])
            self.title_label.setText(t["title_md_to_word"])
            self.files_group.setTitle(t["files_frame_md"])
            self.select_button.setText(t["select_button_md"])
            self.toggle_button.setText(t["toggle_button_md_to_word"])
        else:
            self.setWindowTitle(t["title_word_to_md"])
            self.title_label.setText(t["title_word_to_md"])
            self.files_group.setTitle(t["files_frame_word"])
            self.select_button.setText(t["select_button_word"])
            self.toggle_button.setText(t["toggle_button_word_to_md"])

        # Настройки
        self.settings_group.setTitle(t["settings_frame"])
        self.font_label_widget.setText(t["font_label"])
        self.width_label_widget.setText(t["width_label"])

        # Вкладки
        self.notebook.setTabText(0, t["tab_files"])
        text_tab_idx = self.notebook.indexOf(self.tab_text_widget)
        if text_tab_idx != -1:
            self.notebook.setTabText(text_tab_idx, t["tab_text"])
        self.text_input_label.setText(t["text_input_label"])

        # Язык кнопки
        self.language_button.setText("EN" if self.current_language == "ru" else "RU")

        # Кнопки удаления
        self.remove_button.setText(t["remove_button"])
        self.remove_all_button.setText(t["remove_all_button"])

        # Папка
        self.output_group.setTitle(t["output_frame"])
        if not self.output_directory:
            self.output_label.setText(t["output_label_default"])
        self.output_button.setText(t["output_button"])

        # Статус
        current_status = self.status_label.text()
        ru, en = self.translations["ru"], self.translations["en"]
        if current_status in (ru["status_ready"], en["status_ready"]):
            self.status_label.setText(t["status_ready"])
        elif current_status in (ru["status_finished"], en["status_finished"]):
            self.status_label.setText(t["status_finished"])

        self.convert_button.setText(t["convert_button"])

    # ------------------------------------------------------------------ #
    #  Обработчики настроек
    # ------------------------------------------------------------------ #

    def _on_font_change(self, font_name: str):
        """Обновляет шрифт конвертера при выборе из ComboBox"""
        if self.current_converter_type == "md_to_word":
            self.converter.default_font_name = font_name
            t = self.translations[self.current_language]
            print(t["font_changed"].format(font=font_name))

    def _on_width_change(self, value: int):
        """Обновляет размер шрифта при изменении SpinBox"""
        if self.current_converter_type == "md_to_word":
            self.converter.default_font_size = Pt(value)
            t = self.translations[self.current_language]
            print(t["size_changed"].format(size=value))

    # ------------------------------------------------------------------ #
    #  Работа с файлами
    # ------------------------------------------------------------------ #

    def _select_files(self):
        """Выбор файлов для конвертации"""
        t = self.translations[self.current_language]
        if self.current_converter_type == "md_to_word":
            filter_str = "Markdown files (*.md);;All files (*.*)"
            title = t["select_button_md"]
        else:
            filter_str = "Word files (*.docx);;All files (*.*)"
            title = t["select_button_word"]

        files, _ = QFileDialog.getOpenFileNames(self, title, "", filter_str)

        for file in files:
            if file not in self.selected_files:
                self.selected_files.append(file)
                self.files_listbox.addItem(os.path.basename(file))

        self.status_label.setText(
            t["files_selected"].format(count=len(self.selected_files))
        )

    def _remove_selected_files(self):
        """Удаление выбранных файлов из списка"""
        selected_rows = sorted(
            [idx.row() for idx in self.files_listbox.selectedIndexes()],
            reverse=True
        )
        for index in selected_rows:
            self.files_listbox.takeItem(index)
            del self.selected_files[index]

        t = self.translations[self.current_language]
        self.status_label.setText(
            t["files_selected"].format(count=len(self.selected_files))
        )

    def _remove_all_files(self):
        """Удаляет все файлы из списка"""
        self.files_listbox.clear()
        self.selected_files.clear()
        t = self.translations[self.current_language]
        self.status_label.setText(t["files_selected"].format(count=0))

    def _select_output_directory(self):
        """Выбор папки для сохранения"""
        t = self.translations[self.current_language]
        directory = QFileDialog.getExistingDirectory(self, t["output_button"])
        if directory:
            self.output_directory = directory
            self.output_label.setText(directory)

    # ------------------------------------------------------------------ #
    #  Конвертация
    # ------------------------------------------------------------------ #

    def _convert_files(self):
        """Конвертация выбранных файлов или текста"""
        t = self.translations[self.current_language]

        # Определяем активную вкладку: 0 — Файлы, 1 — Текст
        current_tab = self.notebook.currentIndex()

        if current_tab == 1:
            # --- Вкладка «Текст» ---
            content = self.text_input.toPlainText().strip()
            if not content:
                QMessageBox.warning(
                    self, "Предупреждение", "Введите текст для конвертации"
                )
                return

            if self.current_converter_type == "md_to_word":
                filter_str = "Word files (*.docx)"
                default_suffix = ".docx"
            else:
                filter_str = "Markdown files (*.md)"
                default_suffix = ".md"

            output_path, _ = QFileDialog.getSaveFileName(
                self, t["save_as_title"], "", filter_str
            )
            if not output_path:
                return

            # Добавляем расширение, если отсутствует
            if not output_path.endswith(default_suffix):
                output_path += default_suffix

            self.status_label.setText(
                t["status_converting"].format(filename="Text")
            )
            QApplication.processEvents()

            if self.current_converter_type == "md_to_word":
                success, message = self.converter.convert_content(
                    content, output_path
                )
            else:
                try:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    success, message = True, "Успешно сохранено"
                except Exception as e:
                    success, message = False, str(e)

            if success:
                QMessageBox.information(
                    self, "Успех",
                    t["success_message_word"].format(count=1)
                )
            else:
                QMessageBox.critical(self, t["error_title"], message)

            self.status_label.setText(t["status_finished"])
            return

        # --- Вкладка «Файлы» ---
        if not self.selected_files:
            QMessageBox.warning(
                self, "Предупреждение", t["warning_no_files"]
            )
            return

        self.progress.setValue(0)
        self.progress.setMaximum(len(self.selected_files))

        success_count = 0
        errors = []

        for i, input_file in enumerate(self.selected_files):
            filename = os.path.basename(input_file)
            self.status_label.setText(
                t["status_converting"].format(filename=filename)
            )
            QApplication.processEvents()

            input_path_obj = Path(input_file)
            if self.current_converter_type == "md_to_word":
                output_filename = input_path_obj.stem + ".docx"
            else:
                output_filename = input_path_obj.stem + ".md"

            # Определение пути сохранения
            if self.output_directory:
                output_path = os.path.join(
                    self.output_directory, output_filename
                )
            else:
                output_path = os.path.join(
                    os.path.dirname(input_file), output_filename
                )

            success, message = self.converter.convert_file(
                input_file, output_path
            )

            if success:
                success_count += 1
            else:
                errors.append(f"{filename}: {message}")

            self.progress.setValue(i + 1)
            QApplication.processEvents()

        # Результат
        if errors:
            error_text = "\n".join(errors)
            QMessageBox.warning(
                self,
                t["error_title"],
                t["error_message"].format(
                    success=success_count,
                    errors=len(errors),
                    details=error_text
                )
            )
        else:
            if self.current_converter_type == "md_to_word":
                msg = t["success_message"].format(
                    count=success_count,
                    font=self.converter.default_font_name,
                    size=self.converter.default_font_size.pt
                )
            else:
                msg = t["success_message_word"].format(count=success_count)
            QMessageBox.information(self, "Успех", msg)

        self.status_label.setText(t["status_finished"])


def main():
    """Главная функция"""
    app = QApplication(sys.argv)
    window = ConverterGUI()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
